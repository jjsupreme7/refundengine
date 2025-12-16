"""
Fast Batch Refund Analyzer (Enhanced)

High-quality, efficient invoice analysis using:
- EnhancedRAG: Corrective RAG + reranking + query expansion
- ExcelFileWatcher: Intelligent row-level change tracking
- Smart caching (vendor DB, invoice extraction, RAG results)
- Batch AI analysis (20 items per API call)
- State-aware legal research

Key Features:
- Only analyzes new/changed rows (skips already-analyzed data)
- Database-backed tracking using file and row hashes
- Higher quality legal research with AI-powered reranking
- Maintains fast batch processing for efficiency

INTELLIGENT CHANGE DETECTION (NEW):
- INPUT column changes (new invoice file, new PO, changed amounts) → Re-analyze
- OUTPUT column changes (corrections to AI analysis) → Skip re-analysis (feedback only)
- Uses excel_cell_changes table if file is in versioning system
- Falls back to row-level hash checking for non-versioned files

Usage:
    # Analyze only changed rows
    python analysis/fast_batch_analyzer.py \\
        --excel "Master Refunds.xlsx" --state washington

    # Test mode (force re-analysis of first N rows)
    python analysis/fast_batch_analyzer.py \\
        --excel "Master Refunds.xlsx" --state washington --limit 10
"""

import argparse
import io
import json
import os
import re
import sys
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from typing import Any, Dict, List, Optional

# Fix Windows console encoding for emojis
if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

import pandas as pd
from tqdm import tqdm

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent.parent))

# Imports
from dotenv import load_dotenv  # noqa: E402

from core.database import get_supabase_client  # noqa: E402
from core.enhanced_rag import EnhancedRAG  # noqa: E402
from core.excel_column_definitions import INPUT_COLUMNS, is_input_column  # noqa: E402
from core.excel_file_watcher import ExcelFileWatcher  # noqa: E402
from core.historical_rates import HistoricalRateDB  # noqa: E402
from core.model_router import ModelRouter  # noqa: E402
from core.site_id_lookup import SiteIDLookup  # noqa: E402
from core.wa_tax_rate_lookup import compare_rate, get_correct_rate  # noqa: E402
from scripts.utils.smart_cache import SmartCache  # noqa: E402

# Load environment
load_dotenv()

# Initialize Supabase (optional - will work without it)
try:
    supabase = get_supabase_client()
except ValueError:
    print("⚠️  Supabase not configured - vendor patterns disabled")
    supabase = None

# Initialize multi-model router
# Routes tasks to optimal models:
#   - GPT-5 for extraction/vision (best document understanding)
#   - Claude Sonnet 4.5 for tax analysis & legal reasoning
#   - Claude Opus 4.5 for high-stakes decisions ($25k+)
#   - GPT-4o-mini for validation tasks
router = ModelRouter()

# Known AI hallucination corrections (repealed/invalid → valid)
CITATION_CORRECTIONS = {
    "WAC 458-20-155": "WAC 458-20-15502",  # Repealed 2013 → current software rule
}

# Module-level cache for vendor HQ locations discovered during analysis
_vendor_hq_cache: Dict[str, Dict] = {}


def load_vendor_locations() -> Dict[str, Dict]:
    """Load vendor HQ locations from JSON file, indexed by vendor name."""
    path = Path("knowledge_base/vendors/vendor_locations.json")
    if path.exists():
        with open(path, "r") as f:
            data = json.load(f)
            return {v["vendor_name"].upper(): v for v in data.get("vendors", [])}
    return {}


def save_vendor_location(
    vendor_name: str, hq_city: str, hq_state: str, confidence: int
):
    """Save a newly researched vendor HQ location to the JSON file."""
    from datetime import datetime

    path = Path("knowledge_base/vendors/vendor_locations.json")
    try:
        # Load existing data
        if path.exists():
            with open(path, "r") as f:
                data = json.load(f)
        else:
            data = {
                "enrichment_date": "",
                "total_vendors": 0,
                "statistics": {},
                "vendors": [],
            }

        # Check if vendor already exists (avoid duplicates)
        existing_names = {v["vendor_name"].upper() for v in data.get("vendors", [])}
        if vendor_name.upper() in existing_names:
            return  # Already saved

        # Create new vendor entry
        new_vendor = {
            "headquarters_city": hq_city,
            "headquarters_state": hq_state,
            "headquarters_country": "US",
            "confidence": confidence,
            "reasoning": "Auto-discovered during batch analysis via AI inference.",
            "vendor_name": vendor_name.upper(),
            "source": "batch_analysis_ai",
            "research_date": datetime.now().isoformat(),
        }

        # Append and update counts
        data["vendors"].append(new_vendor)
        data["total_vendors"] = len(data["vendors"])
        data["enrichment_date"] = datetime.now().isoformat()

        # Save back
        with open(path, "w") as f:
            json.dump(data, f, indent=2)

    except Exception as e:
        print(f"  ⚠️  Failed to save vendor location: {e}")


def research_vendor_hq(vendor_name: str) -> Optional[Dict]:
    """Research vendor headquarters location via AI inference.

    Called for unknown vendors not in vendor_locations.json.
    Results are cached in _vendor_hq_cache.
    """
    normalized = vendor_name.upper().strip()

    # Skip if already cached
    if normalized in _vendor_hq_cache:
        return _vendor_hq_cache.get(normalized)

    try:
        prompt = f"""Where is the company "{vendor_name}" headquartered?
Return brief JSON only:
{{"headquarters_city": "city or null", "headquarters_state": "2-letter US state code or null"}}"""

        result = router.execute(task="analysis", prompt=prompt, stakes=0)
        content = result["content"].strip()
        if content.startswith("```"):
            content = content.split("\n", 1)[1].rsplit("```", 1)[0].strip()
        info = json.loads(content)

        hq_city = info.get("headquarters_city")
        hq_state = info.get("headquarters_state")

        if hq_state and hq_state.lower() not in ("null", "unknown", ""):
            _vendor_hq_cache[normalized] = {
                "headquarters_city": hq_city or "",
                "headquarters_state": hq_state.upper(),
                "confidence": 60,
            }
            # Persist to JSON file for future runs
            save_vendor_location(vendor_name, hq_city or "", hq_state.upper(), 60)
            return _vendor_hq_cache[normalized]
    except Exception:
        pass  # Silently fail - HQ lookup is optional

    return None


def validate_citation(citation: str) -> str:
    """Validate citation format and correct known errors.

    Accepts any properly formatted WAC/RCW citation (including rare ones).
    Corrects known hallucinations. Rejects invalid formats.
    """
    citation = citation.strip()

    # Check if it's a known correction
    if citation in CITATION_CORRECTIONS:
        return CITATION_CORRECTIONS[citation]

    # Validate format: WAC xxx-xx-xxxxx with optional subsections like (3)(a)
    wac_pattern = r"^WAC \d{3}-\d{1,2}-\d{3,5}(\([a-zA-Z0-9]+\))*$"
    if re.match(wac_pattern, citation, re.IGNORECASE):
        return citation

    # Validate format: RCW xx.xx.xxx with optional subsections
    rcw_pattern = r"^RCW \d{2}\.\d{2}\.\d{3}(\([a-zA-Z0-9]+\))*$"
    if re.match(rcw_pattern, citation, re.IGNORECASE):
        return citation

    # Invalid format - filter out
    return ""


# Initialize cache
cache = SmartCache()

# Initialize EnhancedRAG (also uses the router internally)
enhanced_rag = EnhancedRAG()


# ==========================================
# PATTERN QUERY FUNCTIONS
# ==========================================


def get_vendor_pattern(vendor_name: str, tax_type: str) -> Optional[Dict]:
    """Get historical vendor pattern from Supabase"""
    if supabase is None:
        return None
    try:
        result = (
            supabase.table("vendor_products")
            .select("*")
            .eq("vendor_name", vendor_name)
            .eq("tax_type", tax_type)
            .maybe_single()
            .execute()
        )
        if result is None:
            return None
        return result.data
    except Exception as e:
        print(f"  ⚠️  Vendor pattern lookup failed for {vendor_name}: {e}")
        return None


def get_refund_basis_patterns(tax_type: str, limit: int = 10) -> List[Dict]:
    """Get top refund basis patterns for tax type"""
    if supabase is None:
        return []
    try:
        result = (
            supabase.table("refund_citations")
            .select("*")
            .eq("tax_type", tax_type)
            .order("usage_count", desc=True)
            .limit(limit)
            .execute()
        )
        if result is None or result.data is None:
            return []
        return result.data
    except Exception as e:
        print(f"  ⚠️  Refund basis pattern lookup failed: {e}")
        return []


def get_keyword_patterns(tax_type: str) -> Dict[str, List[str]]:
    """Get all keyword patterns as dict by category"""
    if supabase is None:
        return {}
    try:
        result = (
            supabase.table("keyword_patterns")
            .select("*")
            .eq("tax_type", tax_type)
            .execute()
        )

        if result is None or result.data is None:
            return {}

        # Convert to dict: {"tax_categories": [...], "product_types": [...]}
        return {row["category"]: row["keywords"] for row in result.data}
    except Exception as e:
        print(f"  ⚠️  Keyword pattern lookup failed: {e}")
        return {}


def get_human_corrections(context: str = None, limit: int = 20) -> List[Dict]:
    """
    Get human corrections from user_feedback table.
    If context is provided, returns corrections similar to the context.
    Otherwise, falls back to recency-based retrieval.

    Args:
        context: Invoice context (vendor + description) to find similar corrections
        limit: Maximum number of corrections to return
    """
    if supabase is None:
        return []

    try:
        # If context provided, use similarity search
        if context:
            from openai import OpenAI

            openai_client = OpenAI()

            # Generate embedding for context
            response = openai_client.embeddings.create(
                input=context[:8000],  # Limit to avoid token overflow
                model="text-embedding-3-small",
            )
            query_embedding = response.data[0].embedding

            # Search by similarity using RPC function
            result = supabase.rpc(
                "search_corrections",
                {
                    "query_embedding": query_embedding,
                    "match_threshold": 0.5,
                    "match_count": limit,
                },
            ).execute()

            if result and result.data:
                return result.data

        # Fallback: recency-based (original behavior)
        result = (
            supabase.table("user_feedback")
            .select("query, suggested_answer, feedback_comment")
            .eq("feedback_type", "correction")
            .order("created_at", desc=True)
            .limit(limit)
            .execute()
        )

        if result is None or result.data is None:
            return []
        return result.data

    except Exception as e:
        print(f"  ⚠️  Human corrections lookup failed: {e}")
        return []


def research_vendor_for_ambiguous(
    vendor_name: str, description: str, tax_type: str = "sales_tax", amount: float = 0
) -> dict:
    """
    Research vendor when description is ambiguous.
    Tier 1: Check database for historical patterns (fast, free)
    Tier 2: AI inference based on vendor name (fast, cheap)

    NOTE: Web search disabled for cost/speed. Enable by setting ENABLE_WEB_SEARCH=true

    Returns dict with vendor context for AI analysis.
    """
    import os

    # Tier 1: Check database first
    db_info = get_vendor_pattern(vendor_name, tax_type)
    if db_info and db_info.get("historical_sample_count", 0) >= 3:
        return {
            "source": "database",
            "business_type": db_info.get("typical_product_type", "Unknown"),
            "context": f"Based on {db_info['historical_sample_count']} previous transactions",
            "typical_refund_basis": db_info.get("typical_refund_basis", "Unknown"),
            "confident": True,
            "sources": [],
        }

    # Tier 2: AI inference (fast, cheap) - skip web search unless explicitly enabled
    if os.getenv("ENABLE_WEB_SEARCH", "").lower() != "true":
        # Use GPT-4o inference instead of web search (much faster and cheaper)
        try:
            fallback_prompt = f"""Based on the vendor name "{vendor_name}" and description "{description}",
what type of business is this likely to be? Where is their headquarters located?
Return brief JSON:
{{"business_type": "type", "typical_services": "description", "headquarters_city": "city or null", "headquarters_state": "2-letter state code or null"}}"""

            result = router.execute(
                task="analysis",
                prompt=fallback_prompt,
                stakes=0,
            )
            content = result["content"].strip()
            if content.startswith("```"):
                content = content.split("\n", 1)[1].rsplit("```", 1)[0].strip()
            vendor_info = json.loads(content)

            # Cache HQ location if found
            hq_city = vendor_info.get("headquarters_city")
            hq_state = vendor_info.get("headquarters_state")
            if hq_state and hq_state.lower() not in ("null", "unknown", ""):
                _vendor_hq_cache[vendor_name.upper()] = {
                    "headquarters_city": hq_city or "",
                    "headquarters_state": hq_state.upper(),
                    "confidence": 60,  # AI-inferred, lower confidence
                }

            return {
                "source": "ai_inference",
                "business_type": vendor_info.get("business_type", "Unknown"),
                "context": vendor_info.get("typical_services", ""),
                "headquarters_city": hq_city,
                "headquarters_state": hq_state,
                "confident": False,
                "sources": [],
            }
        except Exception:
            return {
                "source": "unknown",
                "business_type": "Unknown",
                "context": "",
                "confident": False,
                "sources": [],
            }

    # Tier 3: REAL web search (only if ENABLE_WEB_SEARCH=true)
    try:
        prompt = f"""Research this vendor and their invoice for Washington State tax classification:

Vendor: {vendor_name}
Invoice description: {description}
Tax type: {tax_type}

Please search for:
1. What type of company is {vendor_name}? What services/products do they provide?
2. How should "{description}" be classified for WA state {tax_type}?
3. Are there any relevant WA DOR rules, RCW citations, or WAC references?
4. Is this likely taxable, exempt, or requires more information?
5. Where is {vendor_name} headquartered (city and state)?

Provide your analysis with:
- Business type and typical services
- Headquarters location (city, state)
- WA tax classification recommendation
- Key questions that would affect the classification
- Sources you found (include URLs)

Be thorough - this affects tax refund decisions."""

        print(f"  🔍 Web searching for: {vendor_name}...")
        result = router.execute_with_web_search(
            prompt=prompt,
            stakes=amount,  # Use amount to determine model quality
        )

        # Parse the response - it's free-form analysis, not JSON
        content = result.get("content", "")
        sources = result.get("sources", [])

        # Extract key info from the response
        business_type = "Unknown"
        tax_implications = ""

        # Simple extraction from response
        content_lower = content.lower()
        if "software" in content_lower or "saas" in content_lower:
            business_type = "Software/SaaS"
        elif "consulting" in content_lower or "professional services" in content_lower:
            business_type = "Professional Services"
        elif "construction" in content_lower or "contractor" in content_lower:
            business_type = "Construction Services"
        elif "telecom" in content_lower or "tower" in content_lower:
            business_type = "Telecommunications Infrastructure"
        elif "hardware" in content_lower or "equipment" in content_lower:
            business_type = "Hardware/Equipment"

        # Extract headquarters location from response (simple pattern matching)
        hq_city = None
        hq_state = None
        # Look for patterns like "headquartered in City, ST" or "based in City, State"
        hq_patterns = [
            r"headquartered in ([A-Za-z\s]+),?\s*([A-Z]{2})",
            r"based in ([A-Za-z\s]+),?\s*([A-Z]{2})",
            r"headquarters (?:is |are )?(?:in |located in )?([A-Za-z\s]+),?\s*([A-Z]{2})",
        ]
        for pattern in hq_patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                hq_city = match.group(1).strip()
                hq_state = match.group(2).upper()
                break

        # Cache HQ location if found
        if hq_state:
            _vendor_hq_cache[vendor_name.upper()] = {
                "headquarters_city": hq_city or "",
                "headquarters_state": hq_state,
                "confidence": 75,  # Web search is more confident
            }

        return {
            "source": "web_search",
            "business_type": business_type,
            "context": (
                content[:500] if len(content) > 500 else content
            ),  # First 500 chars
            "full_analysis": content,  # Full analysis for detailed review
            "tax_implications": tax_implications,
            "headquarters_city": hq_city,
            "headquarters_state": hq_state,
            "sources": sources,
            "confident": True,  # Web search is more confident than AI inference
        }

    except Exception as e:
        print(f"  ⚠️  Web search failed for {vendor_name}: {e}")
        # Fallback to basic AI inference if web search fails
        try:
            fallback_prompt = f"""Based on the vendor name "{vendor_name}" and description "{description}",
what type of business is this likely to be? Where is their headquarters located?
Return brief JSON:
{{"business_type": "type", "typical_services": "description", "headquarters_city": "city or null", "headquarters_state": "2-letter state code or null"}}"""

            result = router.execute(
                task="analysis",
                prompt=fallback_prompt,
                stakes=0,
            )
            content = result["content"].strip()
            if content.startswith("```"):
                content = content.split("\n", 1)[1].rsplit("```", 1)[0].strip()
            vendor_info = json.loads(content)

            # Cache HQ location if found
            hq_city = vendor_info.get("headquarters_city")
            hq_state = vendor_info.get("headquarters_state")
            if hq_state and hq_state.lower() not in ("null", "unknown", ""):
                _vendor_hq_cache[vendor_name.upper()] = {
                    "headquarters_city": hq_city or "",
                    "headquarters_state": hq_state.upper(),
                    "confidence": 60,
                }

            return {
                "source": "ai_inference_fallback",
                "business_type": vendor_info.get("business_type", "Unknown"),
                "context": vendor_info.get("typical_services", ""),
                "headquarters_city": hq_city,
                "headquarters_state": hq_state,
                "confident": False,
                "sources": [],
            }
        except Exception:
            return {
                "source": "unknown",
                "business_type": "Unknown",
                "context": "",
                "confident": False,
                "sources": [],
            }


def is_description_ambiguous(description: str) -> bool:
    """
    Check if an invoice description is vague/ambiguous and needs deeper analysis.
    """
    desc_lower = description.lower()

    # Ambiguous keywords that need context to classify
    AMBIGUOUS_KEYWORDS = [
        "inspection",
        "assessment",
        "evaluation",
        "review",
        "consulting",
        "consultation",
        "advisory",
        "services",
        "professional services",
        "service fee",
        "installation",
        "setup",
        "configuration",
        "maintenance",
        "support",
        "management",
        "site",
        "field",
        "on-site",
        "project",
        "engagement",
        "retainer",
        "fee",
        "charge",
        "cost",
    ]

    # Check for ambiguous patterns
    ambiguous_count = sum(1 for kw in AMBIGUOUS_KEYWORDS if kw in desc_lower)

    # If multiple ambiguous keywords or very short description
    if ambiguous_count >= 2:
        return True

    # Very short descriptions are often ambiguous
    if len(description.split()) <= 3:
        return True

    return False


def _extract_pdf_text(file_path: str, max_pages: int = None) -> str:
    """
    Extract text from PDF pages (fast, no API call needed).

    Args:
        file_path: Path to PDF file
        max_pages: If specified, only extract first N pages (for early-stopping)

    This is the preferred method for machine-generated PDFs which have
    embedded text. Falls back to vision API only for scanned documents.
    """
    import pdfplumber

    try:
        with pdfplumber.open(file_path) as pdf:
            if len(pdf.pages) == 0:
                return ""

            pages_to_read = pdf.pages[:max_pages] if max_pages else pdf.pages
            all_text = []
            for i, page in enumerate(pages_to_read):
                text = page.extract_text() or ""
                if text.strip():
                    all_text.append(f"[Page {i+1}]\n{text}")

            return "\n\n".join(all_text)
    except Exception as e:
        print(f"  ⚠️ PDF text extraction failed: {e}")
        return ""


def _has_sufficient_invoice_data(extracted_data: dict) -> bool:
    """
    Check if we have enough invoice data to skip remaining pages.

    Used for early-stopping optimization - if first 2 pages have the key
    data, no need to read pages 3-10 of a long contract.

    Required: vendor_name, invoice_number
    Nice to have: total_amount OR line_items (need at least one)
    """
    if not extracted_data or "error" in extracted_data:
        return False

    # Must have vendor and invoice number
    vendor = extracted_data.get("vendor_name")
    invoice_num = extracted_data.get("invoice_number")
    if not vendor or vendor == "Unknown" or not invoice_num or invoice_num == "Unknown":
        return False

    # Should have at least one of: total or line items
    has_total = (
        extracted_data.get("total_amount") and extracted_data.get("total_amount") > 0
    )
    has_items = (
        extracted_data.get("line_items")
        and len(extracted_data.get("line_items", [])) > 0
    )

    return has_total or has_items


def _extract_pdf_all_pages_images(file_path: str, max_pages: int = 5) -> list:
    """
    Extract first N pages of PDF as base64 images for vision API.

    Used as fallback for scanned PDFs where text extraction fails.
    Limits to max_pages to avoid sending huge documents to API.
    """
    import base64
    from io import BytesIO

    import pdfplumber

    images = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages[:max_pages]):
                img = page.to_image(resolution=150)
                buffered = BytesIO()
                img.original.save(buffered, format="PNG")
                images.append(base64.b64encode(buffered.getvalue()).decode("utf-8"))
        return images
    except Exception as e:
        print(f"  ⚠️ PDF image extraction failed: {e}")
        return []


def _extract_pdf_image(file_path: str) -> str:
    """Extract first page of PDF as base64 image (legacy, use _extract_pdf_all_pages_images instead)"""
    import base64
    from io import BytesIO

    import pdfplumber

    with pdfplumber.open(file_path) as pdf:
        if len(pdf.pages) == 0:
            return None

        first_page = pdf.pages[0]
        img = first_page.to_image(resolution=150)

        buffered = BytesIO()
        img.original.save(buffered, format="PNG")
        return base64.b64encode(buffered.getvalue()).decode("utf-8")


def _extract_image(file_path: str) -> str:
    """Extract JPG/PNG image as base64"""
    import base64

    with open(file_path, "rb") as f:
        return base64.b64encode(f.read()).decode("utf-8")


def _extract_excel_invoice(file_path: str) -> Dict[str, Any]:
    """
    Extract invoice data from Excel file

    Expected columns: Vendor, Invoice Number, Date, Line Item, Amount, Tax
    """
    try:
        df = pd.read_excel(file_path)

        # Try to find common column patterns
        vendor_col = next((col for col in df.columns if "vendor" in col.lower()), None)
        invoice_col = next(
            (
                col
                for col in df.columns
                if "invoice" in col.lower() and "number" in col.lower()
            ),
            None,
        )
        date_col = next((col for col in df.columns if "date" in col.lower()), None)
        amount_col = next(
            (
                col
                for col in df.columns
                if "amount" in col.lower() or "total" in col.lower()
            ),
            None,
        )
        tax_col = next((col for col in df.columns if "tax" in col.lower()), None)
        desc_col = next(
            (
                col
                for col in df.columns
                if "description" in col.lower()
                or "item" in col.lower()
                or "product" in col.lower()
            ),
            None,
        )
        # Location columns for rate lookup
        ship_city_col = next(
            (
                col
                for col in df.columns
                if "ship" in col.lower() and "city" in col.lower()
            ),
            None,
        )
        ship_state_col = next(
            (
                col
                for col in df.columns
                if "ship" in col.lower() and "state" in col.lower()
            ),
            None,
        )
        ship_zip_col = next(
            (
                col
                for col in df.columns
                if "ship" in col.lower() and "zip" in col.lower()
            ),
            None,
        )
        subtotal_col = next(
            (col for col in df.columns if "subtotal" in col.lower()), None
        )

        # Extract line items
        line_items = []
        for idx, row in df.iterrows():
            line_items.append(
                {
                    "description": (
                        row.get(desc_col, "Unknown") if desc_col else "Unknown"
                    ),
                    "amount": float(row.get(amount_col, 0)) if amount_col else 0,
                    "tax": float(row.get(tax_col, 0)) if tax_col else 0,
                }
            )

        return {
            "vendor_name": (
                df[vendor_col].iloc[0] if vendor_col and len(df) > 0 else "Unknown"
            ),
            "invoice_number": (
                df[invoice_col].iloc[0] if invoice_col and len(df) > 0 else "Unknown"
            ),
            "invoice_date": (
                str(df[date_col].iloc[0]) if date_col and len(df) > 0 else None
            ),
            "total_amount": sum(item["amount"] for item in line_items),
            "tax_amount": sum(item["tax"] for item in line_items),
            "line_items": line_items,
            # Location data for rate lookup
            "ship_to_city": (
                str(df[ship_city_col].iloc[0]) if ship_city_col and len(df) > 0 else ""
            ),
            "ship_to_state": (
                str(df[ship_state_col].iloc[0])
                if ship_state_col and len(df) > 0
                else ""
            ),
            "ship_to_zip": (
                str(df[ship_zip_col].iloc[0]) if ship_zip_col and len(df) > 0 else ""
            ),
            "subtotal": (
                float(df[subtotal_col].iloc[0])
                if subtotal_col and len(df) > 0
                else sum(item["amount"] for item in line_items)
            ),
        }
    except Exception as e:
        return {"error": f"Excel extraction failed: {str(e)}"}


def _extract_word_invoice(file_path: str) -> Dict[str, Any]:
    """
    Extract invoice data from Word document

    Uses python-docx to parse tables, then GPT-5 (via router) for text extraction.
    GPT-5 excels at document understanding and structured data extraction.
    """
    try:
        from docx import Document

        doc = Document(file_path)

        # Extract all text
        full_text = "\n".join([para.text for para in doc.paragraphs])

        # Try to extract tables (invoices often have tables)
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                tables_text.append(row_text)

        combined_text = full_text + "\n\n" + "\n".join(tables_text)

        # Use router for extraction task (routes to GPT-5)
        prompt = f"""Extract invoice data from this Word document text and return JSON:

{combined_text}

Return format:
{{
  "vendor_name": "Company Name",
  "invoice_number": "INV-123",
  "invoice_date": "2024-01-15",
  "total_amount": 100.00,
  "tax_amount": 10.00,
  "line_items": [
    {{
      "description": "Product/service description",
      "amount": 100.00,
      "tax": 10.00
    }}
  ]
}}"""

        result = router.execute(
            task="extraction",
            prompt=prompt,
            stakes=0,  # Extraction is a preliminary task
        )

        return json.loads(result["content"])

    except ImportError:
        return {"error": "python-docx not installed. Run: pip install python-docx"}
    except Exception as e:
        return {"error": f"Word extraction failed: {str(e)}"}


def extract_invoice_with_vision(file_path: str) -> Dict[str, Any]:
    """
    Extract invoice data using GPT-5 Vision (via router) or structured parsing.

    GPT-5 has improved visual perception and 45% fewer hallucinations than GPT-4o,
    making it ideal for document extraction tasks.

    Supports:
    - PDF files (GPT-5 Vision)
    - Images: JPG, PNG (GPT-5 Vision)
    - Excel files: .xlsx, .xls (pandas parsing)
    - Word docs: .docx (python-docx parsing)

    With caching support
    """
    # Check cache first
    if cached := cache.get_invoice_data(file_path):
        return cached

    print(f"  📄 Extracting: {Path(file_path).name}")

    file_ext = Path(file_path).suffix.lower()

    try:
        import base64  # noqa: E402
        from io import BytesIO  # noqa: E402

        # Route to appropriate extraction method based on file type
        if file_ext in [".xlsx", ".xls"]:
            return _extract_excel_invoice(file_path)
        elif file_ext == ".docx":
            return _extract_word_invoice(file_path)

        # For PDFs: Try text extraction first (fast)
        # Uses early-stopping: read first 2 pages, only get more if needed
        if file_ext == ".pdf":
            # EARLY STOPPING: Try first 2 pages first
            pdf_text_partial = _extract_pdf_text(file_path, max_pages=2)
            if len(pdf_text_partial) > 100:
                # Good text extraction - try parsing first 2 pages
                prompt = f"""Extract invoice data from this document text and return JSON:

DOCUMENT TEXT:
{pdf_text_partial[:8000]}

Return this JSON format:
{{
  "vendor_name": "Company Name",
  "invoice_number": "INV-123",
  "invoice_date": "2024-01-15",
  "total_amount": 100.00,
  "tax_amount": 10.00,
  "site_id": "Site/location ID if present (patterns like SE01403, 9ME080A, SA03261A, LIB3401A)",
  "line_items": [
    {{"description": "Product/service description", "amount": 100.00, "tax": 10.00}}
  ]
}}"""
                result = router.execute(
                    task="extraction",
                    prompt=prompt,
                    stakes=0,
                )

                # Check if we got sufficient data from first 2 pages
                try:
                    content = result["content"].strip()
                    if content.startswith("```"):
                        first_newline = content.find("\n")
                        if first_newline != -1:
                            content = content[first_newline + 1 :]
                        if content.endswith("```"):
                            content = content[:-3].strip()
                    partial_data = json.loads(content)

                    if _has_sufficient_invoice_data(partial_data):
                        # Early stop - we have enough data!
                        print(f"    (early stop: found key data in first 2 pages)")
                        cache.set_invoice_data(file_path, partial_data)
                        return partial_data
                except (json.JSONDecodeError, KeyError):
                    pass  # Failed to parse, will try full document

                # Need more data - extract ALL pages
                pdf_text_full = _extract_pdf_text(file_path)
                if len(pdf_text_full) > len(pdf_text_partial):
                    print(f"    (reading full document for more data)")
                    prompt = f"""Extract invoice data from this document text and return JSON:

DOCUMENT TEXT:
{pdf_text_full[:8000]}

Return this JSON format:
{{
  "vendor_name": "Company Name",
  "invoice_number": "INV-123",
  "invoice_date": "2024-01-15",
  "total_amount": 100.00,
  "tax_amount": 10.00,
  "site_id": "Site/location ID if present (patterns like SE01403, 9ME080A, SA03261A, LIB3401A)",
  "line_items": [
    {{"description": "Product/service description", "amount": 100.00, "tax": 10.00}}
  ]
}}"""
                    result = router.execute(
                        task="extraction",
                        prompt=prompt,
                        stakes=0,
                    )
            else:
                # Scanned PDF - use vision with ALL pages (up to 5)
                print(f"    (scanned PDF, using vision for {Path(file_path).name})")
                images = _extract_pdf_all_pages_images(file_path, max_pages=5)
                if not images:
                    return {"error": "Failed to extract images from scanned PDF"}

                prompt = """Extract invoice data from these document pages and return JSON:
{
  "vendor_name": "Company Name",
  "invoice_number": "INV-123",
  "invoice_date": "2024-01-15",
  "total_amount": 100.00,
  "tax_amount": 10.00,
  "site_id": "Site/location ID if present (patterns like SE01403, 9ME080A, SA03261A, LIB3401A)",
  "line_items": [
    {"description": "Product/service description", "amount": 100.00, "tax": 10.00}
  ]
}"""
                result = router.execute(
                    task="extraction",
                    prompt=prompt,
                    stakes=0,
                    images=images,  # Pass ALL page images
                )

        elif file_ext in [".jpg", ".jpeg", ".png"]:
            img_base64 = _extract_image(file_path)
            if not img_base64:
                return {"error": "Failed to extract image from file"}

            prompt = """Extract invoice data and return JSON:
{
  "vendor_name": "Company Name",
  "invoice_number": "INV-123",
  "invoice_date": "2024-01-15",
  "total_amount": 100.00,
  "tax_amount": 10.00,
  "site_id": "Site/location ID if present (patterns like SE01403, 9ME080A, SA03261A, LIB3401A)",
  "line_items": [
    {"description": "Product/service description", "amount": 100.00, "tax": 10.00}
  ]
}"""
            result = router.execute(
                task="extraction",
                prompt=prompt,
                stakes=0,
                images=[img_base64],
            )
        else:
            return {"error": f"Unsupported file format: {file_ext}"}

        # Strip markdown code fences if present (GPT often wraps JSON in ```json...```)
        content = result["content"].strip()
        if content.startswith("```"):
            # Remove opening fence (```json or ```)
            first_newline = content.find("\n")
            if first_newline != -1:
                content = content[first_newline + 1 :]
            # Remove closing fence
            if content.endswith("```"):
                content = content[:-3].strip()

        try:
            invoice_data = json.loads(content)
        except json.JSONDecodeError as json_err:
            print(f"  ⚠️ Invalid JSON from invoice extraction: {json_err}")
            print(f"     Content preview: {content[:200]}...")
            return {"error": "JSON parse failed", "raw_preview": content[:500]}

        # Cache result
        cache.set_invoice_data(file_path, invoice_data)

        return invoice_data

    except Exception as e:
        print(f"  ❌ Error: {e}")
        return {"error": str(e)}


def categorize_product(description: str, vendor_info: Optional[Dict] = None) -> str:
    """
    Categorize product based on description and vendor info.
    Uses comprehensive keyword lists for accurate hardware vs software detection.
    """
    desc_lower = description.lower()

    # Comprehensive hardware keywords
    HARDWARE_KEYWORDS = [
        # Network equipment
        "server",
        "switch",
        "router",
        "firewall",
        "modem",
        "gateway",
        "access point",
        "wireless",
        "ethernet",
        "nic",
        "network card",
        # Telecom equipment
        "radio",
        "antenna",
        "transceiver",
        "amplifier",
        "baseband",
        "cell site",
        "tower",
        "microwave",
        "rf ",
        "lte",
        "5g equipment",
        # Computing hardware
        "laptop",
        "desktop",
        "workstation",
        "computer",
        "pc ",
        "monitor",
        "display",
        "keyboard",
        "mouse",
        "headset",
        "printer",
        "scanner",
        "projector",
        "camera",
        # Storage/components
        "storage",
        "disk",
        "drive",
        "ssd",
        "hdd",
        "tape",
        "memory",
        "ram",
        "dimm",
        "cpu",
        "processor",
        "blade",
        "chassis",
        "enclosure",
        "rack",
        "cabinet",
        # Power/infrastructure
        "ups",
        "battery",
        "power supply",
        "pdu",
        "generator",
        "cooling",
        "hvac",
        "air conditioning",
        # Cables/accessories
        "cable",
        "fiber",
        "connector",
        "adapter",
        "converter",
        "mount",
        "bracket",
        "tray",
        "shelf",
        # General hardware terms
        "hardware",
        "equipment",
        "device",
        "appliance",
        "unit",
        "physical",
        "tangible",
    ]

    # Software/digital keywords
    SOFTWARE_KEYWORDS = [
        # Software licenses
        "license",
        "software",
        "perpetual",
        "subscription",
        # Specific vendors/products
        "vmware",
        "microsoft 365",
        "office 365",
        "adobe",
        "oracle",
        "sap",
        "salesforce",
        "servicenow",
        "workday",
        "splunk",
        "citrix",
        "red hat",
        "linux",
        "windows server",
        # Cloud/SaaS
        "saas",
        "cloud",
        "azure",
        "aws",
        "google cloud",
        "gcp",
        "workspace",
        "online",
        "hosted",
        "web-based",
        # Security software
        "antivirus",
        "endpoint protection",
        "security suite",
        "firewall software",
        "vpn software",
        # Digital goods
        "digital",
        "electronic delivery",
        "download",
    ]

    # Professional services keywords
    SERVICES_KEYWORDS = [
        "consulting",
        "professional services",
        "advisory",
        "hours",
        "labor",
        "installation",
        "configuration",
        "implementation",
        "support",
        "maintenance",
        "training",
        "assessment",
    ]

    # Cloud infrastructure keywords
    CLOUD_INFRA_KEYWORDS = [
        "ec2",
        "vm",
        "virtual machine",
        "iaas",
        "paas",
        "compute instance",
        "cloud server",
        "virtual server",
    ]

    # Check in order: software first (MPU potential), then hardware, then services
    if any(kw in desc_lower for kw in SOFTWARE_KEYWORDS):
        # Double-check it's not hardware with "software" in name
        if not any(
            hw in desc_lower
            for hw in ["hardware", "server", "switch", "router", "printer"]
        ):
            return "software_license"

    if any(kw in desc_lower for kw in CLOUD_INFRA_KEYWORDS):
        return "iaas_paas"

    if any(
        kw in desc_lower
        for kw in ["saas", "subscription", "365", "workspace", "cloud service"]
    ):
        return "saas_subscription"

    if any(kw in desc_lower for kw in SERVICES_KEYWORDS):
        return "professional_services"

    if any(kw in desc_lower for kw in HARDWARE_KEYWORDS):
        return "tangible_personal_property"

    # Default to "other" - but the AI prompt now has rules to handle this
    return "other"


def search_legal_docs(category: str, state: str = "WA") -> List[Dict]:
    """
    Search legal documents for category using EnhancedRAG (with caching)

    Uses corrective RAG + reranking for higher quality results.
    """
    # Check cache
    if cached := cache.get_rag_results(category, state):
        return cached

    print(f"  📚 Searching legal docs for: {category}")

    try:
        # Create query based on category
        category_queries = {
            "saas_subscription": (
                "Washington tax law cloud services SaaS "
                "multi-point use digital products WAC 458-20-15502"
            ),
            "professional_services": (
                "Washington tax law professional services "
                "primarily human effort consulting exemption"
            ),
            "iaas_paas": (
                "Washington tax law cloud infrastructure "
                "IaaS PaaS virtual machine allocation"
            ),
            "software_license": (
                "Washington tax law software licenses prewritten custom "
                "multiple points of use MPU WAC 458-20-15503"
            ),
            "tangible_personal_property": (
                "Washington tax tangible personal property hardware "
                "equipment taxable no exemption"
            ),
            # Handle "other" with general tax queries
            "other": (
                "Washington sales tax exemptions tangible personal property "
                "hardware software services taxability"
            ),
        }

        query_text = category_queries.get(category, category_queries["other"])

        # Force database retrieval to get actual law text (not just structured rules)
        # The structured rules in tax_rules.json don't have full legal text,
        # which causes the AI to hallucinate explanations for citations.
        result = enhanced_rag.search_with_decision(
            query=query_text,
            context={
                "product_type": category,
            },
            force_retrieval=True,  # Always get real law text from database
        )

        # Extract results from decision response
        legal_docs = result.get("results", [])

        # Log the decision made by agentic RAG
        action = result.get("action", "UNKNOWN")
        cost_saved = result.get("cost_saved", 0)
        confidence = result.get("confidence", 0)
        print(
            f"    🤖 Agentic RAG decision: {action} (confidence: {confidence:.2f}, saved: ${cost_saved:.4f})"
        )

        # Cache results
        cache.set_rag_results(category, legal_docs, state)

        return legal_docs

    except Exception as e:
        print(f"  ⚠️  RAG search failed: {e}")
        return []


def repair_json(text: str) -> str:
    """
    Attempt to repair common JSON errors from LLM output.
    Handles: single quotes, trailing commas, unquoted keys, truncated responses.
    """
    import re

    # Replace single quotes with double quotes (careful not to break contractions)
    # Only replace single quotes that look like JSON delimiters
    text = re.sub(r"(?<=[{,:\[])\s*'([^']+)'\s*(?=[},:}\]])", r'"\1"', text)

    # Remove trailing commas before } or ]
    text = re.sub(r",\s*([}\]])", r"\1", text)

    # Try to find valid JSON object
    start = text.find("{")
    if start != -1:
        # Find matching closing brace
        brace_count = 0
        for i, c in enumerate(text[start:], start):
            if c == "{":
                brace_count += 1
            elif c == "}":
                brace_count -= 1
                if brace_count == 0:
                    return text[start : i + 1]

    return text


# Structured Output Schema - guarantees valid JSON from OpenAI
ANALYSIS_RESPONSE_SCHEMA = {
    "type": "json_schema",
    "json_schema": {
        "name": "tax_analysis_response",
        "strict": True,
        "schema": {
            "type": "object",
            "properties": {
                "analyses": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "transaction_id": {"type": "integer"},
                            "product_classification": {"type": "string"},
                            "refund_basis": {"type": "string"},
                            "legal_citations": {
                                "type": "array",
                                "items": {"type": "string"},
                            },
                            "confidence_score": {"type": "integer"},
                            "estimated_refund_amount": {"type": "number"},
                            "explanation": {"type": "string"},
                            "needs_review": {"type": "boolean"},
                            "follow_up_questions": {
                                "type": "array",
                                "items": {"type": "string"},
                            },
                        },
                        "required": [
                            "transaction_id",
                            "product_classification",
                            "refund_basis",
                            "legal_citations",
                            "confidence_score",
                            "estimated_refund_amount",
                            "explanation",
                            "needs_review",
                            "follow_up_questions",
                        ],
                        "additionalProperties": False,
                    },
                }
            },
            "required": ["analyses"],
            "additionalProperties": False,
        },
    },
}


def analyze_batch(
    items: List[Dict],
    legal_context: Dict[str, List],
    state: str = "WA",
    tax_type: str = "sales_tax",
    is_retry: bool = False,
) -> List[Dict]:
    """
    Batch analyze multiple line items at once
    Now with historical pattern intelligence!
    """
    # Get unique vendors in this batch
    vendors_in_batch = list(set(item["vendor"] for item in items))

    # Query vendor patterns for this batch
    vendor_patterns = {}
    for vendor_name in vendors_in_batch:
        pattern = get_vendor_pattern(vendor_name, tax_type)
        if pattern:
            vendor_patterns[vendor_name] = pattern

    # Query top refund basis patterns for this tax type
    refund_basis_patterns = get_refund_basis_patterns(tax_type, limit=10)

    # Build prompt for batch analysis
    items_text = ""
    vendor_reasoning = {}  # Store reasoning per item index for AI_Reasoning column
    for i, item in enumerate(items, 1):
        vendor_info = item.get("vendor_info", {})
        line_item = item.get("line_item", {})
        category = item.get("category", "other")
        vendor_name = item["vendor"]

        # Use excel_description if available (more reliable), fallback to line_item
        product_desc = item.get("excel_description") or line_item.get(
            "description", "Unknown"
        )

        items_text += f"\n{i}. Vendor: {vendor_name}\n"
        # [INVOICE DATA] section - data extracted from invoice document
        items_text += f"   [INVOICE DATA] Product: {product_desc}\n"
        items_text += f"   [INVOICE DATA] Amount: ${item['amount']:,.2f}\n"
        items_text += f"   [INVOICE DATA] Tax Charged: ${item['tax']:,.2f}\n"
        items_text += f"   Category: {category}\n"

        # Location and rate info (for Wrong Rate detection)
        ship_city = item.get("ship_to_city", "")
        ship_state = item.get("ship_to_state", "")
        ship_zip = item.get("ship_to_zip", "")
        inv_date = item.get("invoice_date", "")
        rate_charged = item.get("tax_rate_charged", 0)
        correct_rate = item.get("correct_rate")
        rate_diff = item.get("rate_difference")
        if ship_city or ship_state:
            items_text += (
                f"   [INVOICE DATA] Ship-To: {ship_city}, {ship_state} {ship_zip}\n"
            )
            items_text += f"   [INVOICE DATA] Invoice Date: {inv_date}\n"
            items_text += f"   [INVOICE DATA] Tax Rate Charged: {rate_charged}%\n"
            if correct_rate:
                items_text += (
                    f"   [WA DOR RATE] Official Rate for {ship_city}: {correct_rate}%\n"
                )
                if rate_diff and rate_diff > 0.1:
                    items_text += f"   [WA DOR RATE] *** OVERCHARGE: {rate_diff}% above official rate ***\n"
                elif rate_diff and rate_diff < -0.1:
                    items_text += f"   [WA DOR RATE] (Undercharge: {abs(rate_diff)}% below official rate)\n"

        # Add all Excel columns (especially important for use tax with Tax_Remitted)
        excel_row_data = item.get("excel_row_data", {})
        if excel_row_data:
            # Show key Excel columns that might not be in invoice
            key_cols = [
                "Tax_Remitted",
                "Tax_Rate",
                "Tax_Rate_Charged",
                "Account",
                "GL_Code",
                "Notes",
            ]
            shown_cols = []
            for col in key_cols:
                for excel_col, val in excel_row_data.items():
                    if col.lower() in excel_col.lower() and val:
                        items_text += f"   [EXCEL DATA] {excel_col}: {val}\n"
                        shown_cols.append(excel_col)
            # Show any other columns not already displayed
            for col, val in excel_row_data.items():
                if col not in shown_cols and col not in [
                    "Vendor",
                    "Description",
                    "Amount",
                    "Tax_Amount",
                    "Tax_Paid",
                ]:
                    # Skip columns we already show elsewhere
                    if val and len(str(val)) < 100:  # Skip very long values
                        items_text += f"   [EXCEL DATA] {col}: {val}\n"

        # Add PO context if available
        po_content = item.get("po_content", {})
        if po_content and not po_content.get("error"):
            po_desc = ""
            po_line_items = po_content.get("line_items", [])
            if po_line_items:
                po_desc = ", ".join(
                    [
                        li.get("description", "")
                        for li in po_line_items[:3]
                        if li.get("description")
                    ]
                )
            if po_desc:
                items_text += f"   [PO DATA] Line Items: {po_desc[:200]}\n"

        # Add vendor pattern info if available
        if vendor_name in vendor_patterns:
            vp = vendor_patterns[vendor_name]
            items_text += f"   Historical Pattern: {vp.get('historical_sample_count', 0)} transactions, "
            items_text += (
                f"typical refund: {vp.get('typical_refund_basis', 'Unknown')}\n"
            )

        if vendor_info:
            items_text += (
                f"   Vendor Type: {vendor_info.get('business_model', 'Unknown')}\n"
            )

        # Check for ambiguous descriptions and research vendor if needed
        if is_description_ambiguous(product_desc):
            items_text += f"   [AMBIGUOUS] *** DESCRIPTION IS VAGUE - REQUIRES DEEPER ANALYSIS ***\n"
            # Research vendor for context - pass amount for stakes-based model selection
            vendor_research = research_vendor_for_ambiguous(
                vendor_name, product_desc, tax_type, amount=item["amount"]
            )
            if (
                vendor_research.get("business_type")
                and vendor_research["business_type"] != "Unknown"
            ):
                items_text += f"   [VENDOR RESEARCH] Business Type: {vendor_research['business_type']}\n"
                if vendor_research.get("context"):
                    items_text += (
                        f"   [VENDOR RESEARCH] Context: {vendor_research['context']}\n"
                    )
                if vendor_research.get("tax_implications"):
                    items_text += f"   [VENDOR RESEARCH] Tax Notes: {vendor_research['tax_implications']}\n"
                items_text += f"   [VENDOR RESEARCH] Source: {vendor_research.get('source', 'unknown')}\n"
                # Include web search sources if available
                if vendor_research.get("sources"):
                    items_text += f"   [VENDOR RESEARCH] URLs: {', '.join(vendor_research['sources'][:3])}\n"
                # Store full analysis for AI_Reasoning column
                if vendor_research.get("full_analysis"):
                    vendor_reasoning[i] = vendor_research["full_analysis"]

    # Build legal context - INCLUDE ACTUAL DOCUMENT CONTENT
    legal_text = ""
    for category, docs in legal_context.items():
        if docs:
            legal_text += f"\n[LEGAL CONTEXT] {category.upper()} - Relevant Laws:\n"
            for doc in docs[:3]:  # Top 3 per category
                doc_title = doc.get("document_title", "Unknown")
                doc_citation = doc.get("citation", "N/A")
                # Get the actual document content - check multiple possible field names
                doc_content = (
                    doc.get("content")
                    or doc.get("text")
                    or doc.get("chunk_text")
                    or doc.get("page_content")
                    or ""
                )

                legal_text += (
                    f"\n  [LEGAL CONTEXT] === {doc_title} ({doc_citation}) ===\n"
                )
                if doc_content:
                    # Limit to 2000 chars per doc to avoid token overflow
                    content_preview = doc_content[:2000]
                    if len(doc_content) > 2000:
                        content_preview += "... [truncated]"
                    legal_text += f"  [LEGAL CONTEXT] {content_preview}\n"
                else:
                    legal_text += f"  [LEGAL CONTEXT] [No content available]\n"

    # Debug: Show content status
    docs_with_content = sum(
        1
        for cat, docs in legal_context.items()
        for doc in docs[:3]
        if doc.get("content")
        or doc.get("text")
        or doc.get("chunk_text")
        or doc.get("page_content")
    )
    docs_total = sum(min(len(docs), 3) for docs in legal_context.values() if docs)
    print(f"  📄 Legal docs: {docs_with_content}/{docs_total} have content")

    # Build pattern context
    pattern_text = ""
    if refund_basis_patterns:
        pattern_text += (
            f"\nTOP REFUND BASIS PATTERNS ({tax_type.upper().replace('_', ' ')}):\n"
        )
        for pattern in refund_basis_patterns[:5]:
            refund_basis = pattern.get("refund_basis", "Unknown")
            usage_count = pattern.get("usage_count", 0)
            percentage = pattern.get("percentage", 0)
            pattern_text += (
                f"  - {refund_basis}: {usage_count} uses ({percentage:.1f}%)\n"
            )

    # Build context from items for similarity-based correction search
    correction_context = "\n".join(
        [
            f"Vendor: {item.get('vendor_name', '')}, Product: {item.get('description', '')}"
            for item in items[:5]  # Use first 5 items as representative sample
        ]
    )

    # Get human corrections similar to current batch (or fallback to recent)
    human_corrections = get_human_corrections(context=correction_context, limit=15)
    corrections_text = ""
    if human_corrections:
        corrections_text = "\nHUMAN REVIEWER CORRECTIONS (apply these learnings):\n"
        for correction in human_corrections[:10]:
            query = correction.get("query", "")
            answer = correction.get("suggested_answer", "")
            comment = correction.get("feedback_comment", "")
            if query and answer:
                corrections_text += f"  - {query}: Correct answer is '{answer}'\n"

    prompt = f"""You are a Washington State tax refund expert analyzing {tax_type.replace('_', ' ')} transactions.

Analyze these {len(items)} transactions for refund eligibility.

PRIMARY SOURCE - WASHINGTON TAX LAW (80-90% weight):
{legal_text}

SECONDARY SOURCE - HISTORICAL PATTERNS (10-20% weight):
{pattern_text}
{corrections_text}
DECISION HIERARCHY:
1. Base recommendations primarily on Washington State tax law (WAC/RCW)
2. Use historical patterns to:
   - Confirm your legal interpretation (high confidence when they align)
   - Flag potential inconsistencies (review needed when they conflict)
   - Provide context for vendor-specific behavior
3. If law and patterns conflict, law takes precedence (but note the discrepancy)
4. Always cite legal authority (WAC/RCW) when available

CRITICAL - REFUND BASIS RULES:

MPU (Multiple Points of Use) - ONLY applies to:
  - Software licenses (VMware, Microsoft Office, Adobe, SAP, Oracle)
  - Digital goods delivered electronically
  - SaaS/cloud software subscriptions
  - Remotely accessed services/software

MPU does NOT apply to:
  - Hardware (servers, printers, switches, routers, radios, phones, antennas)
  - Tangible personal property (anything physical you can touch)
  - Installation/labor services
  - Consulting/professional services

DECISION TREE:
1. Is it SOFTWARE/DIGITAL? -> Consider "MPU" if used at multiple locations
2. Is it HARDWARE/PHYSICAL? -> "No Refund" (taxable TPP in Washington)
3. Is it SERVICES? -> Check if "primarily human effort" for potential exemption

LOCATION-BASED REFUND OPPORTUNITIES:

Out-of-State (OOS):
  - If ship_to_state is NOT "WA", the sale may qualify for OOS refund
  - Hardware shipped out of state is not subject to WA sales tax
  - Use refund_basis: "OOS" for out-of-state shipments

Wrong Rate (Overcharge):
  - WA tax rates vary by location (city/county/ZIP)
  - We lookup official WA DOR rates for each location - look for "Correct WA Rate" in transaction data
  - If "*** OVERCHARGE ***" is shown, the vendor charged MORE than the official WA rate
  - Use refund_basis: "Wrong Rate" for any transaction with a verified overcharge
  - The difference shown is the recoverable tax overpayment

IMPORTANT - BE DECISIVE:

You MUST make a clear classification and decision for each transaction. Do NOT use:
- "Ambiguous"
- "Requires Review"
- "Unknown"
- "Insufficient Information"

Instead, use the available information to make the BEST classification:
1. Look at the EXTRACTED invoice/PO data first - these have actual product descriptions
2. Use vendor research to understand what the company sells
3. Apply Washington tax law to classify the product/service
4. Make a decision and calculate the refund

If description is vague but vendor type is clear (e.g., software company):
- Classify based on vendor's typical products
- Set confidence to 60-70% (reasonable confidence)
- Explain your reasoning

If truly unclear, default to the MOST LIKELY classification based on:
- Vendor name and industry
- Amount (large amounts suggest software licenses, small amounts suggest services)
- Any keywords in description

TRANSACTIONS:
{items_text}

For EACH transaction, analyze:
1. Product classification - Be SPECIFIC (e.g., "Software License - Enterprise", "SaaS Subscription", "Hardware - Server")
2. Is it taxable in Washington?
3. Refund basis (MPU, Non-taxable, OOS, Wrong Rate, No Refund)
4. Legal citations
5. Confidence (0-100%) - Use 70%+ for clear cases, 50-70% for reasonable inferences
6. Estimated refund amount - CALCULATE THIS:
   - For MPU: refund = tax_amount * 0.80 (80% of tax)
   - For Non-taxable services: refund = tax_amount * 1.00 (100% of tax)
   - For Wrong Rate: refund = difference between charged and correct rate
   - For No Refund: refund = 0
7. Explanation with sources
8. Flag for human review if uncertain
9. Provide specific follow-up questions when flagged

UNCERTAINTY FLAGGING:
Even when making your best determination, flag transactions that need human review.

Set needs_review: true when:
- Confidence is below 70%
- Description is vague (e.g., "Professional Services" with no detail)
- Classification could change based on missing information
- Amount seems unusual for the product type
- You cannot find specific legal text supporting your conclusion
- The transaction type is unclear (service vs. tangible property)
- Ship-to address and Site ID point to different locations (tax rate may differ)
  - For installation services: Site ID location matters (where work is performed)
  - For tangible personal property: Ship-to address matters (where goods delivered)
  - Flag when mismatch exists and classification affects which location applies
- Mixed/bundled transactions (software + hardware + services combined)
- High dollar amount ($50k+ tax) - warrants extra scrutiny regardless of confidence
- Lease vs. purchase ambiguity (different tax treatment)
- Subscription vs. perpetual license unclear (affects taxability and MPU eligibility)
- Installation or labor components mixed with product purchases
- Unusual vendor/product combination (e.g., hardware vendor selling services)
- Out-of-state delivery or nexus questions
- First-time vendor with unusual transaction pattern

Provide follow_up_questions when flagging:
- Be SPECIFIC about what information would change the answer
- Examples: "Is this a perpetual license or a maintenance renewal?"
- Examples: "Was this equipment installed out-of-state?"
- Examples: "Is this a SaaS subscription or on-premise software?"
- Examples: "Was this installation service performed on real property?"
- Examples: "Does the Site ID location or Ship-to address determine where services were performed?"
- Examples: "Is this a lease/rental or an outright purchase?"
- Examples: "Are the bundled services separately stated on the invoice?"
- REQUIRED: When needs_review is true, you MUST provide at least one specific follow_up_question explaining why
- Use empty array [] ONLY when needs_review is false

IMPORTANT: Never flag needs_review: true with an empty follow_up_questions array.
If you're uncertain enough to flag for review, you MUST articulate what specific
information or verification would resolve the uncertainty. Users need actionable
questions, not just a "needs review" flag with no explanation.

Always make your best determination AND flag what to verify.

Return JSON array:
{{
  "analyses": [
    {{
      "transaction_id": 1,
      "product_classification": "...",
      "is_taxable_in_wa": true/false,
      "refund_basis": "<basis>",  // Common: MPU, Non-taxable, OOS, Wrong Rate, No Refund. Use any appropriate legal basis.
      "legal_citations": ["WAC 458-20-15502"],
      "confidence_score": 85,
      "estimated_refund_amount": 4250.00,
      "refund_percentage": 85,
      "explanation": "...",
      "needs_review": false,
      "follow_up_questions": []
    }}
  ]
}}"""

    try:
        system_msg = """You are a Washington State tax refund expert analyzing invoices and purchase orders.

YOUR ANALYSIS MUST:
1. **Apply Washington State tax law** (WAC/RCW) as primary authority
2. **Reference historical patterns** for consistency and confidence
3. **Prioritize legal correctness** over historical precedent
4. **Flag discrepancies** when your analysis differs from typical patterns

CRITICAL - LEGAL CITATIONS REQUIRED:
Every tax decision MUST cite the specific applicable law. Do NOT guess or make up citations.
Think carefully: What WAC or RCW section governs this type of transaction?
- For software/digital products: WAC 458-20-15501, 15502, 15503
- For services (B&O classification): WAC 458-20-224
- For sourcing/location: WAC 458-20-145
- For use tax: WAC 458-20-178, RCW 82.12.020
- For retail sales: RCW 82.08.020
- For construction/installation: WAC 458-20-170, WAC 458-20-171

If you're unsure which law applies, reason through it step by step:
1. What type of product/service is this?
2. Is it tangible or intangible?
3. Is it software, a service, or a digital good?
4. What does Washington law say about that category?

Never skip the citation - every tax determination has a legal basis.

ANTI-HALLUCINATION WARNING:
- ONLY cite what the law actually says - DO NOT fabricate legal language
- If you cite a WAC/RCW, only describe provisions you're certain exist
- When uncertain about specific legal language, say: "Based on general WA tax principles..."
- It's better to say "I'm uncertain about the specific legal provision" than to invent text
- For services: WAC 458-20-224 covers B&O tax CLASSIFICATION, not sales tax exemptions
- Installation services on real property are generally taxable as retail sales (WAC 458-20-170)
- If you cannot find specific legal support in the LEGAL CONTEXT section, flag for review

KNOWLEDGE SOURCES:
- Tax Law Context: Legal citations and rules (PRIMARY - 80-90% weight)
- Historical Patterns: Vendor behavior, refund basis precedent (SECONDARY - 10-20% weight)

When analyzing each transaction:
- Determine legal taxability FIRST (based on WAC/RCW)
- Check historical patterns SECOND (for consistency)
- If they align: High confidence
- If they conflict: Flag for review, recommend legal answer

Provide accurate analysis with legal citations and confidence scores."""

        # Calculate stakes for this batch (sum of tax amounts)
        batch_stakes = sum(item.get("tax", 0) for item in items)

        # Use router for analysis task with structured outputs (guarantees valid JSON)
        # For high-stakes batches ($25k+), router automatically escalates to premium models
        result = router.execute(
            task="analysis",
            prompt=prompt,
            system_prompt=system_msg,
            stakes=batch_stakes,
            temperature=0.2,
            response_format=ANALYSIS_RESPONSE_SCHEMA,  # Guarantees valid JSON
        )

        # Strip markdown code fences if present (Claude often wraps JSON in ```json...```)
        content = result["content"].strip()
        if content.startswith("```"):
            # Remove opening fence (```json or ```)
            first_newline = content.find("\n")
            if first_newline != -1:
                content = content[first_newline + 1 :]
            # Remove closing fence
            if content.endswith("```"):
                content = content[:-3].strip()

        # Extract just the JSON object (AI sometimes adds explanatory text after)
        # Find the first { and the matching closing }
        start_idx = content.find("{")
        if start_idx != -1:
            # Count braces to find matching close
            brace_count = 0
            end_idx = start_idx
            for i, char in enumerate(content[start_idx:], start_idx):
                if char == "{":
                    brace_count += 1
                elif char == "}":
                    brace_count -= 1
                    if brace_count == 0:
                        end_idx = i
                        break
            content = content[start_idx : end_idx + 1]

        # Try to parse JSON with repair and retry on failure
        try:
            parsed = json.loads(content)
        except json.JSONDecodeError as json_err:
            # Try to repair the JSON
            print(f"  ⚠️  JSON parse error: {json_err}")
            print(f"  🔧 Attempting JSON repair...")
            repaired = repair_json(content)
            try:
                parsed = json.loads(repaired)
                print(f"  ✓ JSON repair successful")
            except json.JSONDecodeError:
                # Repair failed - retry the batch if this is first attempt
                if not is_retry:
                    print(f"  🔄 Retrying batch...")
                    return analyze_batch(
                        items, legal_context, state, tax_type, is_retry=True
                    )
                else:
                    print(f"  ❌ JSON repair failed, skipping batch")
                    raise

        print(f"    📊 Model used: {result['model']} (stakes: ${batch_stakes:,.0f})")

        # Add vendor research reasoning to each analysis (for AI_Reasoning column)
        analyses = parsed.get("analyses", [])
        for analysis in analyses:
            tid = analysis.get("transaction_id")
            if tid and tid in vendor_reasoning:
                analysis["ai_reasoning"] = vendor_reasoning[tid]
        return analyses

    except Exception as e:
        print(f"  ❌ Batch analysis error: {e}")
        # Log the problematic content for debugging
        if "content" in dir():
            print(
                f"  📄 Response preview: {content[:200]}..."
                if len(content) > 200
                else f"  📄 Response: {content}"
            )
        # Return empty list so we can continue with other batches
        return []


def get_rows_with_input_changes(file_id: str, hours_lookback: int = 24) -> set:
    """
    Query excel_cell_changes table to find rows where INPUT columns changed.

    This identifies rows that need re-analysis because new data was added
    (e.g., new invoice file, new PO, changed tax amount).

    OUTPUT column changes (like corrections to Analysis_Notes) are ignored
    since those are feedback, not new data requiring re-analysis.

    Args:
        file_id: UUID of the Excel file in excel_file_tracking
        hours_lookback: How far back to check for changes (default: 24 hours)

    Returns:
        Set of row indices that have INPUT column changes
    """
    from datetime import datetime, timedelta

    # Skip if Supabase is not configured
    if supabase is None:
        return set()

    try:
        # Calculate cutoff time
        cutoff_time = (datetime.utcnow() - timedelta(hours=hours_lookback)).isoformat()

        # Query recent cell changes for this file
        result = (
            supabase.table("excel_cell_changes")
            .select("row_index, column_name, old_value, new_value, changed_at")
            .eq("file_id", file_id)
            .gte("changed_at", cutoff_time)
            .execute()
        )

        if not result.data:
            return set()

        # Filter to INPUT column changes only
        rows_needing_reanalysis = set()
        input_changes_count = 0
        output_changes_count = 0

        for change in result.data:
            column_name = change["column_name"]
            row_index = change["row_index"]

            if is_input_column(column_name):
                rows_needing_reanalysis.add(row_index)
                input_changes_count += 1
            else:
                output_changes_count += 1

        print(f"  📊 Cell changes in last {hours_lookback} hours:")
        print(f"     INPUT columns (triggers re-analysis): {input_changes_count}")
        print(f"     OUTPUT columns (feedback only): {output_changes_count}")
        print(f"  📝 Rows needing re-analysis: {len(rows_needing_reanalysis)}")

        return rows_needing_reanalysis

    except Exception as e:
        print(f"  ⚠️  Could not check excel_cell_changes: {e}")
        print(f"     Falling back to standard change detection")
        return set()


def main():
    parser = argparse.ArgumentParser(description="Fast Batch Refund Analyzer")
    parser.add_argument("--excel", required=True, help="Path to Excel file")
    parser.add_argument(
        "--state",
        default="washington",
        help="State to analyze for (default: washington)",
    )
    parser.add_argument(
        "--tax-type",
        choices=["sales_tax", "use_tax"],
        default="sales_tax",
        help="Tax type: sales_tax or use_tax (default: sales_tax)",
    )
    parser.add_argument("--limit", type=int, help="Limit to first N rows (testing)")
    parser.add_argument("--output", help="Output file path")
    parser.add_argument("--site-master", help="Path to site ID master sheet (Excel)")
    parser.add_argument(
        "--site-sheet", help="Sheet name in site master (default: first sheet)"
    )
    parser.add_argument(
        "--rates-folder", help="Path to folder with historical rate files"
    )

    args = parser.parse_args()

    print("🚀 FAST BATCH REFUND ANALYZER")
    print("=" * 70)
    print(f"Excel: {args.excel}")
    print(f"State: {args.state.upper()}")
    print(f"Tax Type: {args.tax_type.replace('_', ' ').title()}")
    print("=" * 70)

    # Load Excel
    print("\n📂 Loading Excel...")
    try:
        df = pd.read_excel(args.excel)
    except FileNotFoundError:
        print(f"❌ Excel file not found: {args.excel}")
        sys.exit(1)
    except Exception as e:
        print(f"❌ Failed to load Excel file: {e}")
        sys.exit(1)
    original_row_count = len(df)

    # Load Site ID master sheet (optional)
    site_lookup = None
    if args.site_master:
        try:
            print(f"\n📍 Loading site ID master sheet...")
            site_lookup = SiteIDLookup(args.site_master, sheet_name=args.site_sheet)
        except Exception as e:
            print(f"  ⚠️  Failed to load site master: {e}")

    # Load historical tax rates (optional)
    historical_rates = None
    if args.rates_folder:
        try:
            print(f"\n📊 Loading historical tax rates...")
            historical_rates = HistoricalRateDB(args.rates_folder)
        except Exception as e:
            print(f"  ⚠️  Failed to load historical rates: {e}")

    # Load vendor HQ locations for out-of-state vendor detection
    vendor_locations = load_vendor_locations()
    if vendor_locations:
        print(f"\n📍 Loaded {len(vendor_locations)} vendor HQ locations")

    # Initialize ExcelFileWatcher for intelligent row tracking
    print("\n🔍 Checking for changes...")
    watcher = ExcelFileWatcher()

    # Try to find file_id from versioning system first
    file_id = None
    try:
        # Check if this file is tracked in the versioning system
        abs_path = str(Path(args.excel).resolve())
        file_name = Path(args.excel).name

        # Try by file_path first
        result = (
            supabase.table("excel_file_tracking")
            .select("id")
            .eq("file_path", abs_path)
            .limit(1)
            .execute()
        )

        # If not found, try by file_name
        if not result or not result.data:
            result = (
                supabase.table("excel_file_tracking")
                .select("id")
                .eq("file_name", file_name)
                .limit(1)
                .execute()
            )

        if result and result.data:
            file_id = result.data[0]["id"]
            print(f"  ✓ Found file in versioning system (ID: {file_id[:8]}...)")
    except Exception as e:
        print(f"  ⚠️  Could not check versioning system: {e}")

    # Get changed rows (new or modified rows only)
    if not args.limit:
        # Option 1: Use versioning system if available (more intelligent)
        if file_id:
            print(f"  🔍 Checking INPUT column changes (versioning-aware)...")
            rows_with_input_changes = get_rows_with_input_changes(
                file_id, hours_lookback=72
            )

            if rows_with_input_changes:
                # Filter DataFrame to only rows with INPUT column changes
                df = df.loc[df.index.isin(rows_with_input_changes)]
                print(
                    f"✓ Found {len(df)} rows with INPUT column changes out of {original_row_count} total"
                )
                print(
                    f"  Skipping {original_row_count - len(df)} rows (no INPUT changes)"
                )
            else:
                print("✓ No INPUT column changes detected!")
                print(
                    "  💡 OUTPUT column changes (corrections) don't trigger re-analysis"
                )
                print("  💡 Use --limit flag to force re-analysis for testing")
                sys.exit(0)

        # Option 2: Fallback to row-level hash checking (less intelligent)
        else:
            print(f"  🔍 Checking row-level changes (hash-based)...")
            changed_rows = watcher.get_changed_rows(args.excel, df)

            if changed_rows:
                changed_indices = [idx for idx, _, _ in changed_rows]
                df = df.loc[changed_indices]
                print(
                    f"✓ Found {len(changed_rows)} changed/new rows out of {original_row_count} total"
                )
                print(
                    f"  Skipping {original_row_count - len(changed_rows)} already-analyzed rows"
                )
            else:
                print("✓ No changes detected - all rows already analyzed!")
                print("\n💡 Tip: Use --limit flag to force re-analysis for testing")
                sys.exit(0)
    else:
        df = df.head(args.limit)
        print(
            f"✓ Test mode: Processing first {args.limit} rows (change tracking disabled)"
        )

    print(f"✓ Analyzing {len(df)} rows")

    # Column name mapping (Excel columns may vary)
    # Try multiple possible column names for flexibility
    def get_col(row, names, default=None):
        """Get value from row, trying multiple possible column names"""
        for name in names if isinstance(names, list) else [names]:
            if name in row.index and pd.notna(row.get(name)):
                return row.get(name)
        return default

    def safe_float(value, default=0.0):
        """
        Safely convert value to float, handling Excel quirks like:
        - Numbers stored as text (green triangle)
        - Currency formatting ($1,234.56)
        - Commas as thousands separators
        - Empty strings or None
        """
        if value is None or value == "":
            return default
        if isinstance(value, (int, float)):
            return float(value)
        if isinstance(value, str):
            # Remove currency symbols, commas, spaces
            cleaned = value.replace("$", "").replace(",", "").replace(" ", "").strip()
            try:
                return float(cleaned)
            except ValueError:
                return default
        try:
            return float(value)
        except (ValueError, TypeError):
            return default

    # Column name variants to support different file formats
    INV_COL_NAMES = [
        "Inv 1 File Name",
        "Inv-1 File Name",  # Denodo/Use Tax actual names
        "Inv-1 FileName",
        "Inv_1_File",
        "Invoice_File_Name_1",
    ]
    INV2_COL_NAMES = [
        "Inv 2 File Name",
        "Inv-2 File Name",  # Denodo/Use Tax actual names
        "Inv-2 FileName",
        "Inv_2_File",
        "Invoice_File_Name_2",
    ]
    PO_COL_NAMES = [
        "PO File Name",  # Denodo/Use Tax actual name
        "PO_FileName",
        "PO_File",
        "Purchase_Order_File_Name",
    ]
    AMOUNT_COL_NAMES = ["Initial Amount", "Amount", "Total_Amount"]
    TAX_COL_NAMES = [
        "hwste_tax_amount_lc",
        "Total Tax",  # Denodo/Use Tax actual names
        "Tax Paid",
        "Tax Remitted",
        "Tax",
        "Tax_Amount",
    ]
    DESC_COL_NAMES = [
        "txz01_po_description",
        "Description",  # Denodo/Use Tax actual names
        "Line_Item_Description",
        "Product_Description",
    ]
    INV_FOLDER_COL_NAMES = [
        "Invoice Folder Path",  # Denodo/Use Tax actual name
        "Invoice_Folder",
        "Inv_Folder",
        "Invoice_Path",
    ]
    PO_FOLDER_COL_NAMES = [
        "PO Folder Path",  # Denodo/Use Tax actual name
        "PO_Folder",
        "Purchase_Order_Folder",
        "PO_Path",
    ]
    SITE_ID_COL_NAMES = ["Site_ID", "SiteID", "Site ID", "Location_ID", "LocationID"]
    # Tax jurisdiction state columns (for out-of-state detection)
    TAX_STATE_COL_NAMES = [
        "tax_jurisdiction_state",
        "sales_tax_state",  # Denodo actual names
        "Tax_State",
        "Jurisdiction_State",
        "State",
    ]
    # KOM analysis/comments columns (skip rows already analyzed)
    KOM_COMMENTS_COL_NAMES = [
        "KOM Analysis & Notes",  # Use Tax actual name
        "Recon Analysis",  # Denodo actual name
        "KOM_Analysis",
        "Analysis_Notes",
        "Comments",
    ]
    # Essential output columns (slim output - only keep these from source + AI outputs)
    ESSENTIAL_OUTPUT_COLS = [
        # Key identifiers
        "Vendor",
        "name1_po_vendor_name",
        "Inv 1 File Name",
        "Inv-1 File Name",
        "Inv 2 File Name",
        "Inv-2 File Name",
        "Invoice Folder Path",
        # Financial
        "Initial Amount",
        "Amount",
        "hwste_tax_amount_lc",
        "Total Tax",
        "Tax Paid",
        "Tax Remitted",
        # Description
        "txz01_po_description",
        "Description",
        # Location
        "tax_jurisdiction_state",
        "sales_tax_state",
        # AI Analysis outputs (added by this script)
        "Product_Desc",
        "Product_Type",
        "Refund_Basis",
        "Citation",
        "Confidence",
        "Estimated_Refund",
        "Explanation",
        "Needs_Review",
        "Follow_Up_Questions",
        "AI_Reasoning",
    ]

    # Get unique invoices - try multiple column names
    inv_col = next((c for c in INV_COL_NAMES if c in df.columns), None)
    inv2_col = next((c for c in INV2_COL_NAMES if c in df.columns), None)
    po_col = next((c for c in PO_COL_NAMES if c in df.columns), None)
    inv_folder_col = next((c for c in INV_FOLDER_COL_NAMES if c in df.columns), None)
    site_id_col = next((c for c in SITE_ID_COL_NAMES if c in df.columns), None)
    if site_id_col and site_lookup:
        print(f"📍 Using Site_ID column: {site_id_col}")
    po_folder_col = next((c for c in PO_FOLDER_COL_NAMES if c in df.columns), None)
    tax_state_col = next((c for c in TAX_STATE_COL_NAMES if c in df.columns), None)
    kom_col = next((c for c in KOM_COMMENTS_COL_NAMES if c in df.columns), None)

    if tax_state_col:
        print(f"📍 Tax jurisdiction column found: {tax_state_col}")
    if kom_col:
        print(f"📝 KOM analysis column found: {kom_col}")
        # Filter to rows without existing KOM analysis (unreviewed rows)
        has_kom = df[kom_col].notna() & (df[kom_col].astype(str).str.strip() != "")
        rows_with_kom = has_kom.sum()
        if rows_with_kom > 0:
            original_count = len(df)
            df = df[~has_kom]
            print(f"   Filtered out {rows_with_kom} rows with existing KOM analysis")
            print(f"   Remaining: {len(df)} unreviewed rows (from {original_count})")

    if inv_col:
        invoice_files = df[inv_col].dropna().unique()
        print(f"\n📄 Found {len(invoice_files)} unique invoices (column: {inv_col})")
    else:
        invoice_files = []
        print("\n⚠️ No invoice column found")

    # Get unique Inv 2 files (supplemental invoices)
    if inv2_col:
        invoice2_files = df[inv2_col].dropna().unique()
        # Exclude files already in inv1 list to avoid duplicate extraction
        invoice2_files = [f for f in invoice2_files if f not in invoice_files]
        if len(invoice2_files) > 0:
            print(
                f"📄 Found {len(invoice2_files)} unique Inv 2 files (column: {inv2_col})"
            )
    else:
        invoice2_files = []

    # Get unique POs
    if po_col:
        po_files = df[po_col].dropna().unique()
        print(f"📋 Found {len(po_files)} unique POs (column: {po_col})")
    else:
        po_files = []
        print("⚠️ No PO column found")

    # Helper function to resolve file paths
    def resolve_file_path(
        filename: str, search_paths: list, folder_override: str = None
    ) -> Optional[str]:
        """Find file in multiple possible locations. folder_override is checked first."""
        if not filename or pd.isna(filename):
            return None
        # Check folder override first (from Excel column)
        if folder_override and not pd.isna(folder_override):
            override_path = Path(folder_override) / filename
            if override_path.exists():
                return str(override_path)
        # Then check default search paths
        for base_path in search_paths:
            file_path = Path(base_path) / filename
            if file_path.exists():
                return str(file_path)
        return None

    # Build folder override mappings from Excel columns
    inv_folder_map = {}
    po_folder_map = {}
    if inv_folder_col:
        for _, row in df.iterrows():
            fname = get_col(row, INV_COL_NAMES)
            folder = row.get(inv_folder_col)
            if fname and folder and not pd.isna(folder):
                inv_folder_map[fname] = folder
            # Also map inv2 files to the same folder
            fname2 = get_col(row, INV2_COL_NAMES)
            if fname2 and folder and not pd.isna(folder):
                inv_folder_map[fname2] = folder
        print(f"📁 Using Invoice_Folder column: {inv_folder_col}")
    if po_folder_col:
        for _, row in df.iterrows():
            fname = get_col(row, PO_COL_NAMES)
            folder = row.get(po_folder_col)
            if fname and folder and not pd.isna(folder):
                po_folder_map[fname] = folder
        print(f"📁 Using PO_Folder column: {po_folder_col}")

    # Helper function for parallel extraction
    def extract_file(args):
        """Extract a single file (for parallel processing)"""
        filename, search_paths, folder_override = args
        file_path = resolve_file_path(filename, search_paths, folder_override)
        if file_path:
            return filename, extract_invoice_with_vision(file_path)
        return filename, {"error": "File not found"}

    # Extract invoices (PARALLEL - 10x faster)
    print("\n📄 Extracting invoices (parallel)...")
    invoice_search_paths = [
        "test_data/synthetic",
        "test_data/sales_tax/invoices",
        "test_data/invoices",
        "client_documents/invoices",
    ]
    invoice_data = {}
    if len(invoice_files) > 0:
        with ThreadPoolExecutor(max_workers=10) as executor:
            invoice_args = [
                (f, invoice_search_paths, inv_folder_map.get(f)) for f in invoice_files
            ]
            futures = {
                executor.submit(extract_file, args): args[0] for args in invoice_args
            }
            for future in tqdm(
                as_completed(futures), total=len(futures), desc="Extracting invoices"
            ):
                filename, result = future.result()
                invoice_data[filename] = result

    # Extract Inv 2 files (supplemental invoices - same search paths)
    if len(invoice2_files) > 0:
        print(f"\n📄 Extracting Inv 2 files (parallel)...")
        with ThreadPoolExecutor(max_workers=10) as executor:
            invoice2_args = [
                (f, invoice_search_paths, inv_folder_map.get(f)) for f in invoice2_files
            ]
            futures = {
                executor.submit(extract_file, args): args[0] for args in invoice2_args
            }
            for future in tqdm(
                as_completed(futures), total=len(futures), desc="Extracting Inv 2"
            ):
                filename, result = future.result()
                invoice_data[filename] = result  # Add to same dict

    # Extract POs (PARALLEL - 10x faster)
    print("\n📋 Extracting Purchase Orders (parallel)...")
    po_search_paths = [
        "test_data/synthetic",
        "test_data/sales_tax/purchase_orders",
        "test_data/purchase_orders",
        "client_documents/purchase_orders",
    ]
    po_data = {}
    if len(po_files) > 0:
        with ThreadPoolExecutor(max_workers=10) as executor:
            po_args = [(f, po_search_paths, po_folder_map.get(f)) for f in po_files]
            futures = {executor.submit(extract_file, args): args[0] for args in po_args}
            for future in tqdm(
                as_completed(futures), total=len(futures), desc="Extracting POs"
            ):
                filename, result = future.result()
                po_data[filename] = result

    # Prepare analysis queue with skip tracking
    print("\n🔍 Matching line items...")
    analysis_queue = []
    out_of_state_queue = []  # Rows to flag as OOS (not analyzed under WA law)
    skipped_rows = {
        "no_invoice_file": [],
        "extraction_failed": [],
        "no_line_item_match": [],
        "out_of_state": [],  # Transactions not subject to WA tax
    }

    for idx, row in df.iterrows():
        # Get invoice file using flexible column names
        inv_file = get_col(row, INV_COL_NAMES)
        inv2_file = get_col(row, INV2_COL_NAMES)
        po_file = get_col(row, PO_COL_NAMES)

        # Try invoice 1 first
        inv_data = None
        inv2_data = None
        if inv_file and inv_file in invoice_data:
            inv_data = invoice_data[inv_file]
            if "error" in inv_data:
                inv_data = None  # Try inv2 or PO instead

        # Try invoice 2 (supplemental)
        if inv2_file and inv2_file in invoice_data:
            inv2_data = invoice_data[inv2_file]
            if "error" in inv2_data:
                inv2_data = None
            elif not inv_data:
                # Use inv2 as primary if inv1 failed
                inv_data = inv2_data
                inv_data["_source"] = "Invoice2"

        # Fall back to PO if no valid invoice
        if not inv_data and po_file:
            # Extract PO filename from path if needed
            po_filename = Path(po_file).name if "/" in str(po_file) else po_file
            if po_filename in po_data and "error" not in po_data.get(po_filename, {}):
                inv_data = po_data[po_filename]
                inv_data["_source"] = "PO"  # Mark as PO-sourced

        # If no document available, create placeholder from Excel data
        if not inv_data:
            excel_desc = get_col(row, DESC_COL_NAMES, "")
            if not excel_desc:
                # No description to analyze - truly skip
                skipped_rows["no_invoice_file"].append(idx)
                continue
            # Create minimal data from Excel
            inv_data = {
                "line_items": [
                    {
                        "description": excel_desc,
                        "amount": safe_float(get_col(row, AMOUNT_COL_NAMES, 0)),
                    }
                ],
                "_source": "Excel",  # Mark as Excel-only
            }

        # Get amount using flexible column names (safe_float handles text-as-number, $, commas)
        amount = safe_float(get_col(row, AMOUNT_COL_NAMES, 0))
        tax = safe_float(get_col(row, TAX_COL_NAMES, 0))

        # Get product description from Excel (more reliable than invoice extraction)
        excel_description = get_col(row, DESC_COL_NAMES, "")

        # Match line item by amount
        line_items = inv_data.get("line_items", [])
        matched_item = None
        for item in line_items:
            item_amount = safe_float(item.get("amount", 0))
            if abs(item_amount - amount) < 1.0:
                matched_item = item
                break

        # If no exact match, use first line item or create placeholder
        if not matched_item and line_items:
            matched_item = line_items[0]
        elif not matched_item:
            matched_item = {"description": excel_description, "amount": amount}

        # Use Excel description if available (often more reliable)
        description = (
            excel_description
            if excel_description
            else matched_item.get("description", "Unknown")
        )

        # Enhance description with Inv 2 data if available (provides additional context)
        if inv2_data and inv_data != inv2_data:
            inv2_line_items = inv2_data.get("line_items", [])
            if inv2_line_items:
                inv2_desc = inv2_line_items[0].get("description", "")
                if inv2_desc and inv2_desc not in description:
                    description = f"{description} | Inv2: {inv2_desc[:200]}"

        vendor = row.get("Vendor", "Unknown")
        vendor_info = cache.get_vendor_info(vendor)

        # Use description for categorization
        category = categorize_product(description, vendor_info)

        # Check for out-of-state transactions (not subject to WA tax analysis)
        tax_jurisdiction = get_col(row, TAX_STATE_COL_NAMES, "")
        if tax_jurisdiction:
            tax_jurisdiction = str(tax_jurisdiction).strip().upper()
        # Consider WA if blank (default), "WA", "WASHINGTON", or empty
        is_wa_transaction = not tax_jurisdiction or tax_jurisdiction in (
            "WA",
            "WASHINGTON",
            "",
        )

        if not is_wa_transaction:
            # Out-of-state transaction - add to separate queue for flagging
            out_of_state_queue.append(
                {
                    "excel_row_idx": idx,
                    "tax_jurisdiction": tax_jurisdiction,
                    "vendor": vendor,
                    "description": description,
                    "amount": amount,
                    "tax": tax,
                }
            )
            skipped_rows["out_of_state"].append((idx, tax_jurisdiction))
            continue  # Skip normal WA analysis

        # Get PO data if available (po_file already retrieved above)
        po_filename = Path(po_file).name if po_file and "/" in str(po_file) else po_file
        po_content = po_data.get(po_filename, {}) if po_file else {}

        # Extract location data from invoice (for Wrong Rate detection)
        ship_to_city = inv_data.get("ship_to_city", "")
        ship_to_state = inv_data.get("ship_to_state", "")
        ship_to_zip = inv_data.get("ship_to_zip", "")
        ship_to_county = ""
        invoice_date = inv_data.get("date", row.get("Date", ""))

        # Get Site ID from Excel column OR extracted from invoice PDF
        site_id = get_col(row, SITE_ID_COL_NAMES) if site_id_col else None
        if not site_id and inv_data:
            # Fall back to Site ID extracted from invoice PDF
            site_id = inv_data.get("site_id")
        is_out_of_state = False
        if site_id and site_lookup:
            location = site_lookup.lookup(site_id)
            if location:
                # Override with master sheet data
                ship_to_city = location.get("city", ship_to_city)
                ship_to_state = location.get("state", ship_to_state)
                ship_to_zip = location.get("zip", ship_to_zip)
                ship_to_county = location.get("county", "")
                # Check if out-of-state
                state_upper = (ship_to_state or "").upper().strip()
                if state_upper and state_upper not in ("WA", "WASHINGTON"):
                    is_out_of_state = True

        # Calculate tax rate from invoice data
        subtotal = inv_data.get("subtotal", 0)
        inv_tax = inv_data.get("tax_amount", inv_data.get("tax", tax))
        tax_rate_charged = (
            (inv_tax / subtotal * 100) if subtotal and subtotal > 0 else 0
        )

        # Look up correct WA tax rate (if we have location)
        correct_rate = None
        rate_difference = None
        rate_source = None

        # Parse invoice date for historical lookup
        invoice_date_obj = None
        if invoice_date:
            try:
                from datetime import datetime

                date_str = str(invoice_date)[:10]
                invoice_date_obj = datetime.strptime(date_str, "%Y-%m-%d").date()
            except (ValueError, TypeError):
                pass

        # Look up correct WA tax rate for Wrong Rate detection
        # Skip if out-of-state (no WA tax applies)
        state_upper = (ship_to_state or "").upper().strip()
        is_wa = state_upper in ("WA", "WASHINGTON", "")  # Empty state assumed WA
        if ship_to_city and is_wa and not is_out_of_state:
            # Determine if invoice is "current" (within last 6 months)
            from datetime import date, timedelta

            is_recent_invoice = False
            if invoice_date_obj:
                days_old = (date.today() - invoice_date_obj).days
                is_recent_invoice = days_old < 180  # Within 6 months

            # Get address for API fallback (if site lookup was used)
            ship_to_address = ""
            if site_id and site_lookup:
                loc = site_lookup.lookup(site_id)
                if loc:
                    ship_to_address = loc.get("address", "")

            # Try historical rates first (if loaded and invoice has date)
            # Now includes API fallback for location code
            if historical_rates and invoice_date_obj:
                correct_rate = historical_rates.get_rate(
                    city=ship_to_city,
                    county=ship_to_county,
                    zip_code=ship_to_zip,
                    invoice_date=invoice_date_obj,
                    address=ship_to_address,
                )
                if correct_rate:
                    rate_source = "historical"

            # Only fall back to current API if:
            # 1. No historical rates loaded, OR
            # 2. Invoice is recent (rates unlikely to have changed much)
            if not correct_rate and (not historical_rates or is_recent_invoice):
                correct_rate = get_correct_rate(ship_to_city, ship_to_zip)
                if correct_rate:
                    rate_source = "current"

            # If old invoice and no historical rate found, skip rate comparison
            if not correct_rate and invoice_date_obj and not is_recent_invoice:
                rate_source = "unavailable"  # Flag that we couldn't get accurate rate

            if correct_rate and tax_rate_charged > 0:
                rate_difference = round(tax_rate_charged - correct_rate, 2)

        # Capture all Excel columns for AI context (especially useful for use tax)
        excel_row_data = {
            str(k): str(v)
            for k, v in row.to_dict().items()
            if pd.notna(v) and str(v).strip()
        }

        # Check for vendor HQ location mismatch (OOS vendor claiming WA work)
        vendor_hq_mismatch = None
        normalized_vendor = vendor.upper().strip()
        # Check pre-loaded locations first, then cached results, then research
        hq_info = vendor_locations.get(normalized_vendor) or _vendor_hq_cache.get(
            normalized_vendor
        )
        # If unknown vendor, research their HQ via AI
        if not hq_info:
            hq_info = research_vendor_hq(vendor)
        if hq_info:
            hq_state = (hq_info.get("headquarters_state") or "").upper()
            hq_city = hq_info.get("headquarters_city", "")
            ship_state_upper = (ship_to_state or "").upper()
            confidence = hq_info.get("confidence", 0)
            # Flag if: HQ outside WA, ship-to in WA (or empty), and confidence >= 60
            if (
                hq_state
                and hq_state != "WA"
                and ship_state_upper in ("WA", "WASHINGTON", "")
                and confidence >= 60
            ):
                vendor_hq_mismatch = {"hq_city": hq_city, "hq_state": hq_state}

        analysis_queue.append(
            {
                "excel_row_idx": idx,
                "vendor": vendor,
                "vendor_info": vendor_info or {},
                "line_item": matched_item,
                "excel_description": description,  # Description from Excel
                "po_content": po_content,  # PO data
                "amount": amount,
                "tax": tax,
                "category": category,
                # Location data for Wrong Rate detection
                "site_id": site_id or "",
                "ship_to_city": ship_to_city,
                "ship_to_state": ship_to_state,
                "ship_to_zip": ship_to_zip,
                "ship_to_county": ship_to_county,
                "is_out_of_state": is_out_of_state,
                "invoice_date": str(invoice_date)[:10] if invoice_date else "",
                "tax_rate_charged": round(tax_rate_charged, 2),
                # WA DOR official rate lookup
                "correct_rate": correct_rate,
                "rate_difference": rate_difference,
                "rate_source": rate_source or "",  # "historical" or "current"
                # All Excel columns for AI context
                "excel_row_data": excel_row_data,
                # Vendor HQ mismatch for OOS vendor detection
                "vendor_hq_mismatch": vendor_hq_mismatch,
            }
        )

    # Report matching results
    total_rows = len(df)
    total_skipped = sum(
        (
            len(v)
            if isinstance(v, list) and not isinstance(v[0] if v else None, tuple)
            else len(v)
        )
        for v in skipped_rows.values()
    )

    print("\n📊 MATCHING SUMMARY:")
    print(f"   Total rows: {total_rows}")
    print(f"   ✓ Matched: {len(analysis_queue)}")
    print(f"   ✗ Skipped: {total_skipped}")

    if total_skipped > 0:
        print("\n⚠️  SKIPPED ROWS BREAKDOWN:")
        if skipped_rows["no_invoice_file"]:
            no_inv_count = len(skipped_rows["no_invoice_file"])
            print(f"   • No invoice file: {no_inv_count} rows")
            if no_inv_count <= 10:
                print(f"     Rows: {skipped_rows['no_invoice_file']}")
        if skipped_rows["extraction_failed"]:
            extract_fail_count = len(skipped_rows["extraction_failed"])
            print("   • Invoice extraction failed: " f"{extract_fail_count} rows")
            if extract_fail_count <= 5:
                for idx, error in skipped_rows["extraction_failed"]:
                    print(f"     Row {idx}: {error}")
        if skipped_rows["no_line_item_match"]:
            no_match_count = len(skipped_rows["no_line_item_match"])
            print(f"   • No line item match: {no_match_count} rows")
            if no_match_count <= 10:
                for idx, amount in skipped_rows["no_line_item_match"]:
                    print(f"     Row {idx}: ${amount:,.2f}")
        if skipped_rows["out_of_state"]:
            oos_count = len(skipped_rows["out_of_state"])
            # Count by state
            state_counts = {}
            for idx, state in skipped_rows["out_of_state"]:
                state_counts[state] = state_counts.get(state, 0) + 1
            print(f"   • Out-of-state (not WA): {oos_count} rows")
            for state, count in sorted(state_counts.items(), key=lambda x: -x[1])[:5]:
                print(f"     {state}: {count} rows")

    if total_skipped > len(analysis_queue):
        analyzed_count = len(analysis_queue)
        print(
            f"\n🚨 WARNING: More rows skipped ({total_skipped}) "
            f"than analyzed ({analyzed_count})!"
        )
        print("   You may want to review your invoice files and data.")

    if len(analysis_queue) == 0:
        print("\n❌ ERROR: No rows matched for analysis!")
        print("   Check that your invoice files exist and amounts match.")
        sys.exit(1)

    # Get unique categories and search legal docs
    print("\n📚 Researching legal context...")
    categories = set(item["category"] for item in analysis_queue)
    legal_context = {}

    for category in categories:
        legal_context[category] = search_legal_docs(category, args.state.upper())

    # Batch analysis
    print(f"\n🤖 Analyzing {len(analysis_queue)} items...")
    batch_size = 15  # Larger batches for cost efficiency
    all_analyses = []

    for i in range(0, len(analysis_queue), batch_size):
        batch = analysis_queue[i : i + batch_size]
        batch_num = i // batch_size + 1
        total_batches = (len(analysis_queue) - 1) // batch_size + 1
        print(f"  Batch {batch_num}/{total_batches}")

        # No delay needed for GPT-4o (higher rate limits than Claude)
        analyses = analyze_batch(
            batch, legal_context, args.state.upper(), args.tax_type
        )

        # Re-number transaction IDs with global offset (batches return 1-N but we need global IDs)
        batch_offset = i
        for analysis in analyses:
            if "transaction_id" in analysis:
                analysis["transaction_id"] = analysis["transaction_id"] + batch_offset

        all_analyses.extend(analyses)

    # Write results back to DataFrame
    print("\n📝 Writing results...")

    # Validate analysis count
    if len(all_analyses) != len(analysis_queue):
        analyses_count = len(all_analyses)
        queue_count = len(analysis_queue)
        print(f"⚠️  WARNING: Got {analyses_count} analyses " f"for {queue_count} items")
        print("   Some results may be missing or duplicated!")

    # Create mapping by transaction_id for safe alignment
    analysis_map = {}
    for analysis in all_analyses:
        tid = analysis.get("transaction_id")
        if tid is not None:
            analysis_map[tid] = analysis

    # Write results using transaction_id matching
    matched_count = 0
    for i, queue_item in enumerate(analysis_queue):
        # Try to match by transaction_id (1-indexed in the prompt)
        transaction_id = i + 1
        analysis = analysis_map.get(transaction_id)

        if analysis is None:
            row_idx = queue_item["excel_row_idx"]
            print(
                "⚠️  No analysis found for transaction "
                f"{transaction_id} (row {row_idx})"
            )
            continue

        idx = queue_item["excel_row_idx"]
        line_item = queue_item["line_item"]

        df.loc[idx, "Product_Desc"] = line_item.get("description", "")
        df.loc[idx, "Product_Type"] = analysis.get("product_classification", "")
        df.loc[idx, "Refund_Basis"] = analysis.get("refund_basis", "")
        # Validate citations before saving (filters hallucinations, corrects known errors)
        raw_citations = analysis.get("legal_citations", [])
        valid_citations = [validate_citation(c) for c in raw_citations]
        valid_citations = [c for c in valid_citations if c]  # Remove empty/invalid
        # Fallback if all citations were filtered out
        if not valid_citations:
            valid_citations = ["CITATION NEEDED - Review Required"]
        df.loc[idx, "Citation"] = ", ".join(valid_citations)
        df.loc[idx, "Confidence"] = f"{analysis.get('confidence_score', 0)}%"
        refund_amount = analysis.get("estimated_refund_amount") or 0

        # Rate difference fallback: if AI says no refund but there's a rate difference, use that
        if refund_amount == 0 and "Rate_Difference_Refund" in df.columns:
            rate_diff_refund = df.loc[idx, "Rate_Difference_Refund"]
            if pd.notna(rate_diff_refund) and float(rate_diff_refund) > 0:
                refund_amount = float(rate_diff_refund)
                df.loc[idx, "Refund_Basis"] = "Rate Difference"

        # Non-taxable fallback: if item is non-taxable and they paid tax, refund full tax amount
        refund_basis = df.loc[idx, "Refund_Basis"]
        if refund_amount == 0 and "non-taxable" in str(refund_basis).lower():
            # For use tax, check Tax Remitted column (note: space, not underscore)
            if "Tax Remitted" in df.columns:
                tax_remitted = df.loc[idx, "Tax Remitted"]
                if pd.notna(tax_remitted) and float(tax_remitted) > 0:
                    refund_amount = float(tax_remitted)
            # For sales tax, check Tax Paid column
            elif "Tax Paid" in df.columns:
                tax_paid = df.loc[idx, "Tax Paid"]
                if pd.notna(tax_paid) and float(tax_paid) > 0:
                    refund_amount = float(tax_paid)

        # MPU fallback: if item qualifies for MPU and refund is $0, calculate 80% refund
        if refund_amount == 0 and "mpu" in str(refund_basis).lower():
            # For use tax, check Tax Remitted column
            if "Tax Remitted" in df.columns:
                tax_remitted = df.loc[idx, "Tax Remitted"]
                if pd.notna(tax_remitted) and float(tax_remitted) > 0:
                    refund_amount = float(tax_remitted) * 0.80  # 80% exempt for MPU
            # For sales tax, check Tax Paid column
            elif "Tax Paid" in df.columns:
                tax_paid = df.loc[idx, "Tax Paid"]
                if pd.notna(tax_paid) and float(tax_paid) > 0:
                    refund_amount = float(tax_paid) * 0.80  # 80% exempt for MPU

        df.loc[idx, "Estimated_Refund"] = f"${refund_amount:,.2f}"
        df.loc[idx, "Explanation"] = analysis.get("explanation", "")
        # AI_Reasoning captures vendor research reasoning for audit trail
        if analysis.get("ai_reasoning"):
            df.loc[idx, "AI_Reasoning"] = analysis.get("ai_reasoning", "")
        # Uncertainty flagging for human review
        df.loc[idx, "Needs_Review"] = (
            "Yes" if analysis.get("needs_review", False) else ""
        )
        follow_up = analysis.get("follow_up_questions", [])
        # Add vendor HQ mismatch question if applicable
        hq_mismatch = queue_item.get("vendor_hq_mismatch")
        if hq_mismatch:
            follow_up = list(follow_up)  # Make mutable copy
            follow_up.append(
                f"Vendor HQ is in {hq_mismatch['hq_city']}, {hq_mismatch['hq_state']}. "
                "Invoice claims work in WA. Verify where services were actually performed."
            )
        df.loc[idx, "Follow_Up_Questions"] = "; ".join(follow_up) if follow_up else ""
        # Rate lookup info (from queue_item, not analysis)
        df.loc[idx, "Resolved_City"] = queue_item.get("ship_to_city", "")
        df.loc[idx, "Resolved_County"] = queue_item.get("ship_to_county", "")
        correct_rate = queue_item.get("correct_rate")
        df.loc[idx, "Correct_Rate"] = f"{correct_rate}%" if correct_rate else ""
        df.loc[idx, "Rate_Source"] = queue_item.get("rate_source", "")
        matched_count += 1

    # Write results for out-of-state transactions (not analyzed under WA law)
    for oos_item in out_of_state_queue:
        idx = oos_item["excel_row_idx"]
        state = oos_item["tax_jurisdiction"]
        tax_amount = oos_item["tax"]

        df.loc[idx, "Product_Desc"] = oos_item.get("description", "")
        df.loc[idx, "Product_Type"] = "Out-of-State Transaction"
        df.loc[idx, "Refund_Basis"] = f"OOS - {state}"
        df.loc[idx, "Citation"] = "N/A - Not subject to WA tax"
        df.loc[idx, "Confidence"] = "100%"
        # Full refund since WA tax shouldn't have been charged
        df.loc[idx, "Estimated_Refund"] = f"${tax_amount:,.2f}"
        df.loc[idx, "Explanation"] = (
            f"Transaction in {state}, not subject to WA sales tax. Tax paid to {state} should be reviewed under {state} tax law."
        )
        df.loc[idx, "Needs_Review"] = "Yes"
        df.loc[idx, "Follow_Up_Questions"] = (
            f"Verify tax was paid to {state}, not WA. Review {state} exemption rules."
        )

    if len(out_of_state_queue) > 0:
        print(f"  📍 Flagged {len(out_of_state_queue)} out-of-state transactions")

    if matched_count < len(analysis_queue):
        queue_count = len(analysis_queue)
        print(
            f"⚠️  Only matched {matched_count}/{queue_count} items - "
            "some rows may be incomplete!"
        )

    # Save output
    if not args.output:
        base_name = Path(args.excel).stem
        args.output = str(Path(args.excel).parent / f"{base_name} - Analyzed.xlsx")

    # Slim down columns to essential output only
    output_cols = [c for c in ESSENTIAL_OUTPUT_COLS if c in df.columns]
    original_col_count = len(df.columns)
    df = df[output_cols]
    print(
        f"\n📊 Output columns: {len(output_cols)} (from {original_col_count} source columns)"
    )

    # Add blank correction columns for human review
    review_cols = [
        "Review_Status",
        "Corrected_Product_Type",
        "Corrected_Refund_Basis",
        "Corrected_Citation",
        "Reviewer_Notes",
    ]
    for col in review_cols:
        df[col] = ""
    print(f"📝 Added {len(review_cols)} review columns for corrections")

    # Replace NaN with empty strings so Excel shows blank cells, not "nan"
    df = df.fillna("")

    try:
        df.to_excel(args.output, index=False)
        print(f"\n✅ Results saved to {args.output}")
    except PermissionError:
        # File might be open in another program - try backup location
        backup_path = args.output.replace(".xlsx", "_backup.xlsx")
        print(f"⚠️ Cannot write to {args.output} (file may be open)")
        try:
            df.to_excel(backup_path, index=False)
            print(f"✅ Results saved to backup: {backup_path}")
        except Exception as e:
            print(f"❌ CRITICAL: Failed to save results: {e}")
            print("   Data may be lost! Check disk space and permissions.")
            sys.exit(1)
    except Exception as e:
        print(f"❌ Failed to save results: {e}")
        sys.exit(1)

    # Update tracking database if not in test mode
    if not args.limit:
        print("\n📊 Updating tracking database...")
        watcher.update_file_tracking(args.excel, len(df))

        # Re-read original file to compute consistent hashes (without output columns)
        original_df = pd.read_excel(args.excel)

        # Update row tracking for successfully analyzed rows
        for i, queue_item in enumerate(analysis_queue):
            if i < matched_count:  # Only update successfully analyzed rows
                idx = queue_item["excel_row_idx"]
                if idx in original_df.index:
                    row_series = original_df.loc[idx]
                    row_hash = watcher.get_row_hash(row_series)
                    watcher.update_row_tracking(args.excel, idx, row_hash)

        print("✓ Tracking database updated")

    print("\n✅ ANALYSIS COMPLETE!")
    print("=" * 70)
    print(f"Output: {args.output}")
    print("\n📊 FINAL SUMMARY:")
    print(f"   Total rows in Excel: {total_rows}")
    print(f"   ✓ Analyzed successfully: {matched_count}")
    print(f"   ✗ Skipped: {total_skipped}")
    if total_skipped > 0:
        coverage_pct = (matched_count / total_rows * 100) if total_rows > 0 else 0
        print(f"   Coverage: {coverage_pct:.1f}%")
        if coverage_pct < 80:
            print("\n⚠️  Low coverage! " "Review skipped rows above to improve results.")
    print("\n💡 Open the Excel file to review results!")


if __name__ == "__main__":
    main()
