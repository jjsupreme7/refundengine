#!/usr/bin/env python3
"""
Legal Documents Ingestion for Washington State Tax Refund Engine

Process legal documents with folder hints and AI verification,
create searchable vector database.

Usage:
    python scripts/ingest_legal_docs.py --folder knowledge_base/
    python scripts/ingest_legal_docs.py --folder knowledge_base/statutes/rcw/
"""

import argparse
import sys
import os
import sqlite3
import hashlib
from pathlib import Path
from datetime import datetime
from tqdm import tqdm

# Vector embeddings
import chromadb
from chromadb.config import Settings
from sentence_transformers import SentenceTransformer

# Import our utilities
sys.path.insert(0, str(Path(__file__).parent))
from document_classifier import extract_text_preview, get_file_metadata, detect_document_type
from metadata_extractor import extract_legal_document_metadata

from dotenv import load_dotenv
load_dotenv()

# Global embedding model
embedding_model = None

def get_embedding_model():
    """Lazy load embedding model."""
    global embedding_model
    if embedding_model is None:
        print("🔄 Loading embedding model (sentence-transformers)...")
        embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        print("✅ Embedding model loaded")
    return embedding_model

def extract_full_text(file_path):
    """Extract complete text from document."""
    file_ext = Path(file_path).suffix.lower()

    if file_ext == '.pdf':
        import pdfplumber
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text() or ""
                text += f"\n{page_text}"
        return text.strip(), len(pdf.pages)

    elif file_ext in ['.docx', '.doc']:
        from docx import Document
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text.strip(), len(doc.paragraphs)

    elif file_ext == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        return text.strip(), text.count('\n') + 1

    else:
        # Fallback to preview
        text = extract_text_preview(file_path)
        return text, 1

def chunk_text(text, chunk_size=500, overlap=50):
    """
    Split text into overlapping chunks of approximately chunk_size words.

    Args:
        text: Full text to chunk
        chunk_size: Target words per chunk
        overlap: Overlap words between chunks

    Returns:
        List of text chunks
    """
    words = text.split()
    chunks = []

    start = 0
    while start < len(words):
        end = start + chunk_size
        chunk_words = words[start:end]
        chunk = ' '.join(chunk_words)
        chunks.append(chunk)

        start += (chunk_size - overlap)

    return chunks

def get_folder_hint(file_path, base_folder):
    """
    Determine folder hint from file path structure.

    Examples:
        knowledge_base/statutes/rcw/file.pdf -> 'rcw'
        knowledge_base/guidance/wtd/file.pdf -> 'wtd'
    """
    try:
        rel_path = Path(file_path).relative_to(Path(base_folder))
        parts = rel_path.parts

        # Look for known folder names
        known_types = ['rcw', 'wac', 'wtd', 'eta', 'case_law']
        for part in parts:
            if part.lower() in known_types:
                return part.lower()

        return None
    except:
        return None

def get_database_connection():
    """Get database connection."""
    project_root = Path(__file__).parent.parent
    db_path = project_root / "database" / "refund_engine.db"
    return sqlite3.connect(str(db_path))

def get_chromadb_client():
    """Get ChromaDB client."""
    project_root = Path(__file__).parent.parent
    chroma_path = project_root / "vector_db" / "chroma"

    client = chromadb.PersistentClient(
        path=str(chroma_path),
        settings=Settings(anonymized_telemetry=False)
    )

    return client

def process_legal_document(file_path, base_folder, conn, chroma_collection):
    """
    Process a single legal document through the full pipeline.

    Returns:
        dict with success status and details
    """
    try:
        filename = Path(file_path).name

        # Get folder hint
        folder_hint = get_folder_hint(file_path, base_folder)

        # Get file metadata
        file_meta = get_file_metadata(file_path)

        # Check for duplicates by content hash
        cursor = conn.cursor()
        cursor.execute(
            "SELECT id, citation FROM legal_documents WHERE content_hash = ?",
            (file_meta['content_hash'],)
        )
        existing = cursor.fetchone()
        if existing:
            return {
                'success': False,
                'filename': filename,
                'reason': f"Duplicate (already processed as {existing[1]})"
            }

        # Extract full text
        full_text, page_count = extract_full_text(file_path)

        if not full_text or len(full_text) < 50:
            return {
                'success': False,
                'filename': filename,
                'reason': "Insufficient text extracted"
            }

        # Classify document
        classification = detect_document_type(file_path, folder_hint)

        if classification['confidence'] < 50:
            return {
                'success': False,
                'filename': filename,
                'reason': f"Low classification confidence ({classification['confidence']}%)"
            }

        doc_type = classification['document_type']

        # Extract metadata
        metadata = extract_legal_document_metadata(
            file_path,
            full_text[:6000],
            doc_type
        )

        if metadata.get('confidence_score', 0) < 40:
            print(f"  ⚠️  Low metadata confidence, proceeding anyway...")

        # Insert into legal_documents table
        cursor.execute("""
            INSERT INTO legal_documents (
                document_type, citation, title, document_date, effective_date,
                expiration_date, file_path, file_format, file_size_bytes,
                processed_date, content_hash, raw_extracted_text, is_current
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            doc_type,
            metadata.get('citation'),
            metadata.get('title'),
            metadata.get('document_date'),
            metadata.get('effective_date'),
            metadata.get('expiration_date'),
            str(file_path),
            file_meta['file_format'],
            file_meta['file_size_bytes'],
            datetime.now().isoformat(),
            file_meta['content_hash'],
            full_text[:50000],  # Store first 50k chars
            1
        ))

        document_id = cursor.lastrowid

        # Insert into document_metadata table
        cursor.execute("""
            INSERT INTO document_metadata (
                document_id, topic_tags, industries, key_concepts,
                tax_types, exemption_categories, referenced_statutes
            ) VALUES (?, ?, ?, ?, ?, ?, ?)
        """, (
            document_id,
            ','.join(metadata.get('topic_tags', [])),
            ','.join(metadata.get('industries', [])),
            ','.join(metadata.get('key_concepts', [])),
            ','.join(metadata.get('tax_types', [])),
            ','.join(metadata.get('exemption_categories', [])),
            ','.join(metadata.get('referenced_statutes', []))
        ))

        # Chunk text
        chunks = chunk_text(full_text, chunk_size=500, overlap=50)

        # Insert chunks into database
        for i, chunk in enumerate(chunks):
            cursor.execute("""
                INSERT INTO document_chunks (
                    document_id, chunk_index, chunk_text, page_number
                ) VALUES (?, ?, ?, ?)
            """, (
                document_id,
                i,
                chunk,
                None  # Could calculate page number if needed
            ))

        # Generate embeddings and store in ChromaDB
        model = get_embedding_model()

        # Batch embed chunks for efficiency
        embeddings = model.encode(chunks, show_progress_bar=False)

        # Prepare ChromaDB data
        chunk_ids = [f"doc_{document_id}_chunk_{i}" for i in range(len(chunks))]
        chunk_metadatas = [
            {
                'document_id': document_id,
                'chunk_index': i,
                'citation': metadata.get('citation', 'Unknown'),
                'document_type': doc_type,
                'title': metadata.get('title', '')[:100],
                'topics': ','.join(metadata.get('topic_tags', [])[:5])
            }
            for i in range(len(chunks))
        ]

        # Add to ChromaDB
        chroma_collection.add(
            ids=chunk_ids,
            embeddings=embeddings.tolist(),
            documents=chunks,
            metadatas=chunk_metadatas
        )

        conn.commit()

        return {
            'success': True,
            'filename': filename,
            'document_id': document_id,
            'document_type': doc_type,
            'citation': metadata.get('citation', 'Unknown'),
            'chunks': len(chunks)
        }

    except Exception as e:
        conn.rollback()
        return {
            'success': False,
            'filename': Path(file_path).name,
            'reason': str(e)
        }

def scan_and_process_legal_docs(folder_path):
    """
    Recursively scan folder and process all legal documents.

    Args:
        folder_path: Root folder to scan
    """
    print("\n" + "="*70)
    print("LEGAL DOCUMENTS INGESTION")
    print("="*70)
    print(f"\n📁 Scanning folder: {folder_path}\n")

    # Find all supported files
    supported_extensions = ['.pdf', '.docx', '.doc', '.txt']
    files = []

    for ext in supported_extensions:
        files.extend(Path(folder_path).rglob(f'*{ext}'))

    print(f"📄 Found {len(files)} documents to process\n")

    if len(files) == 0:
        print("⚠️  No documents found. Exiting.")
        return

    # Connect to databases
    conn = get_database_connection()
    chroma_client = get_chromadb_client()
    chroma_collection = chroma_client.get_or_create_collection("legal_knowledge")

    # Process each file
    results = {
        'success': [],
        'failed': []
    }

    print("🔄 Processing documents...\n")

    for file_path in tqdm(files, desc="Processing", unit="doc"):
        result = process_legal_document(
            str(file_path),
            folder_path,
            conn,
            chroma_collection
        )

        if result['success']:
            results['success'].append(result)
        else:
            results['failed'].append(result)

    conn.close()

    # Print summary
    print("\n" + "="*70)
    print("INGESTION SUMMARY")
    print("="*70)

    print(f"\n✅ Successfully processed: {len(results['success'])}")
    if results['success']:
        print("\n  Documents:")
        for r in results['success'][:10]:  # Show first 10
            print(f"  - {r['filename']}")
            print(f"    Type: {r['document_type']}, Citation: {r['citation']}, Chunks: {r['chunks']}")

        if len(results['success']) > 10:
            print(f"  ... and {len(results['success']) - 10} more")

    print(f"\n❌ Failed: {len(results['failed'])}")
    if results['failed']:
        print("\n  Failed documents:")
        for r in results['failed'][:10]:
            print(f"  - {r['filename']}: {r['reason']}")

        if len(results['failed']) > 10:
            print(f"  ... and {len(results['failed']) - 10} more")

    # Calculate statistics
    total_chunks = sum(r.get('chunks', 0) for r in results['success'])
    print(f"\n📊 Statistics:")
    print(f"  Total documents processed: {len(results['success'])}")
    print(f"  Total chunks created: {total_chunks}")
    print(f"  Average chunks per document: {total_chunks / len(results['success']) if results['success'] else 0:.1f}")

    # Save detailed log
    project_root = Path(__file__).parent.parent
    log_path = project_root / "logs" / f"ingestion_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"
    log_path.parent.mkdir(exist_ok=True)

    with open(log_path, 'w') as f:
        f.write("LEGAL DOCUMENTS INGESTION LOG\n")
        f.write(f"Date: {datetime.now().isoformat()}\n")
        f.write(f"Folder: {folder_path}\n\n")

        f.write("SUCCESSFUL:\n")
        for r in results['success']:
            f.write(f"  {r['filename']} - {r['document_type']} - {r['citation']}\n")

        f.write("\nFAILED:\n")
        for r in results['failed']:
            f.write(f"  {r['filename']} - {r['reason']}\n")

    print(f"\n📝 Detailed log saved to: {log_path}")
    print("\n" + "="*70)
    print("✨ Ingestion complete!")
    print("="*70 + "\n")

def main():
    """Main CLI function."""
    parser = argparse.ArgumentParser(
        description="Ingest legal documents into searchable knowledge base"
    )
    parser.add_argument(
        '--folder',
        required=True,
        help="Root folder containing legal documents"
    )

    args = parser.parse_args()

    # Validate folder exists
    if not os.path.exists(args.folder):
        print(f"❌ Error: Folder not found: {args.folder}")
        return 1

    # Process documents
    scan_and_process_legal_docs(args.folder)

    return 0

if __name__ == "__main__":
    sys.exit(main())
